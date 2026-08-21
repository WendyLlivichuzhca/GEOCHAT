# backend/routes/agentes_ia.py
from flask import Blueprint, jsonify, request, redirect
from flask_jwt_extended import jwt_required, get_jwt_identity, decode_token
from main import get_connection, logger, require_admin_role
import os
import re
import requests
import json
import time
from datetime import datetime, timedelta

agentes_ia_blueprint = Blueprint('agentes_ia', __name__)

@agentes_ia_blueprint.route('/api/agentes-ia', methods=['GET'])
@jwt_required()
def get_agentes_ia():
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("""
            SELECT a.*, d.nombre AS dispositivo_nombre, d.numero_telefono AS dispositivo_numero
            FROM agentes_ia a
            LEFT JOIN dispositivos d ON a.dispositivo_id = d.id
            WHERE a.usuario_id = %s
            ORDER BY a.creado_en DESC
        """, (user_id,))
        agents = cursor.fetchall()
        
        # Formatear las fechas a cadenas ISO
        for agent in agents:
            if agent.get('creado_en'):
                agent['creado_en'] = agent['creado_en'].isoformat()
            if agent.get('actualizado_en'):
                agent['actualizado_en'] = agent['actualizado_en'].isoformat()
                
        return jsonify({"success": True, "data": agents})
    except Exception as e:
        logger.exception("Error al listar agentes de IA")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/asesores', methods=['GET'])
@jwt_required()
def get_active_advisors():
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Asegurar que la columna parent_id existe en la tabla usuarios (migración automática rápida)
        cursor.execute("SHOW COLUMNS FROM usuarios LIKE 'parent_id'")
        if not cursor.fetchone():
            cursor.execute("ALTER TABLE usuarios ADD COLUMN parent_id INT(11) DEFAULT NULL")
            conn.commit()
            logger.info("Columna parent_id agregada con éxito a la tabla usuarios.")

        user_id = get_jwt_identity()
        # Excluye colaboradores con rol 'visor' (solo lectura): no pueden responder,
        # así que no tiene sentido ofrecerlos como destino de una transferencia.
        cursor.execute(
            "SELECT nombre FROM usuarios WHERE activo = 1 AND (id = %s OR parent_id = %s) AND (rol IS NULL OR rol != 'visor') ORDER BY nombre ASC",
            (user_id, user_id)
        )
        rows = cursor.fetchall()
        advisors = [r['nombre'] for r in rows if r.get('nombre')]
        if not advisors:
            cursor.execute("SELECT nombre FROM usuarios WHERE id = %s LIMIT 1", (user_id,))
            self_user = cursor.fetchone()
            if self_user and self_user.get('nombre'):
                advisors = [self_user.get('nombre')]
        return jsonify({"success": True, "advisors": advisors})
    except Exception as e:
        logger.error(f"Error al obtener asesores: {e}")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/stats', methods=['GET'])
@jwt_required()
def get_agentes_ia_stats():
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Conteo total
        cursor.execute("SELECT COUNT(*) AS total FROM agentes_ia WHERE usuario_id = %s", (user_id,))
        total = cursor.fetchone()['total']
        
        # Conteo activos
        cursor.execute("SELECT COUNT(*) AS activos FROM agentes_ia WHERE usuario_id = %s AND activo = 1", (user_id,))
        activos = cursor.fetchone()['activos']
        
        # Por ahora la base de conocimiento se calcula como 0 MB o según tamaño de instrucciones/personalidades
        cursor.execute("""
            SELECT SUM(LENGTH(COALESCE(instrucciones, '')) + LENGTH(COALESCE(personalidad, ''))) AS total_size 
            FROM agentes_ia 
            WHERE usuario_id = %s
        """, (user_id,))
        res_size = cursor.fetchone()
        total_size_bytes = res_size['total_size'] if res_size and res_size['total_size'] else 0
        kb_size_mb = round(total_size_bytes / (1024 * 1024), 2)
        
        return jsonify({
            "success": True,
            "total": total,
            "activos": activos,
            "knowledge_base_mb": kb_size_mb
        })
    except Exception as e:
        logger.exception("Error al obtener estadísticas de agentes de IA")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia', methods=['POST'])
@jwt_required()
def create_agente_ia():
    role_err = require_admin_role()
    if role_err:
        return role_err
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    
    dispositivo_id = data.get('dispositivo_id')
    nombre = data.get('nombre')
    modelo = data.get('modelo', 'gpt-4')
    instrucciones = data.get('instrucciones', '')
    personalidad = data.get('personalidad', '')
    activo = 1 if data.get('activo') else 0
    descripcion_negocio = data.get('descripcion_negocio', '')
    industria = data.get('industria', '')
    objetivo = data.get('objetivo', '')
    
    pasos_captura = data.get('pasos_captura')
    skip_existing_data = 1 if data.get('skip_existing_data') else 0
    seguimientos = data.get('seguimientos')
    reglas_transferencia = data.get('reglas_transferencia')
    reglas_etiquetado = data.get('reglas_etiquetado')
    config_comportamiento = data.get('config_comportamiento')
    
    if not nombre:
        return jsonify({"success": False, "message": "nombre es obligatorio"}), 400
        
    # Generar instrucciones y personalidad a medida con IA si se provee descripción del negocio
    if descripcion_negocio and descripcion_negocio.strip():
        try:
            from main import call_llm_api
            generator_prompt = (
                "Eres un experto en ingeniería de prompts para agentes de IA de WhatsApp.\n"
                f"El usuario quiere crear un agente de IA con las siguientes opciones:\n"
                f"- Nombre: {nombre}\n"
                f"- Industria: {industria}\n"
                f"- Objetivo: {objetivo}\n"
                f"- Descripción del negocio: {descripcion_negocio}\n\n"
                "Diseña un prompt de instrucciones (system instructions) altamente optimizado para WhatsApp. "
                "Debe ser claro, profesional, en español y guiar al bot para lograr el objetivo del negocio. "
                "También define un breve perfil de personalidad amigable y profesional.\n\n"
                "Responde ÚNICAMENTE con un objeto JSON válido que contenga las propiedades 'instrucciones' y 'personalidad' (sin bloques de código markdown ni explicaciones adicionales):\n"
                "{\n"
                "  \"instrucciones\": \"...\",\n"
                "  \"personalidad\": \"...\"\n"
                "}"
            )
            openai_key = os.getenv("OPENAI_API_KEY")
            gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
            nvidia_key = os.getenv("NVIDIA_API_KEY")
            
            res_prompt = call_llm_api(generator_prompt, "Creador de Prompts", openai_key, gemini_key, nvidia_key)
            if res_prompt:
                import re
                json_match = re.search(r'\{.*\}', res_prompt.strip(), re.DOTALL)
                if json_match:
                    parsed = json.loads(json_match.group(0))
                    instrucciones = parsed.get("instrucciones") or instrucciones
                    personalidad = parsed.get("personalidad") or personalidad
        except Exception as e_gen:
            logger.error(f"Error generando prompt base con IA: {e_gen}")
            
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Si no se provee dispositivo_id, auto-asignar el primer dispositivo del usuario
        if not dispositivo_id:
            cursor.execute("""
                SELECT id FROM dispositivos WHERE usuario_id = %s ORDER BY id ASC LIMIT 1
            """, (user_id,))
            device_row = cursor.fetchone()
            dispositivo_id = device_row['id'] if device_row else None
        
        # Validar si el plan permite objetivos de IA avanzados
        cursor.execute(
            """
            SELECT p.permite_todos_objetivos_ia
            FROM suscripciones s
            INNER JOIN planes p ON p.id = s.plan_id
            WHERE s.usuario_id = %s
            ORDER BY FIELD(s.estado, 'activa', 'prueba', 'vencida', 'cancelada'), s.fecha_vencimiento DESC, s.id DESC
            LIMIT 1
            """,
            (user_id,)
        )
        plan_row = cursor.fetchone()
        permite_todos_objetivos = bool(plan_row["permite_todos_objetivos_ia"]) if plan_row else False
        
        obj_final = objetivo or 'preguntas_frecuentes'
        if not permite_todos_objetivos and obj_final != 'preguntas_frecuentes':
            return jsonify({"success": False, "message": "Objetivo avanzado no disponible en tu plan actual. Mejora al Plan Advanced para desbloquearlo."}), 400

        # Insertar agente
        cursor.execute("""
            INSERT INTO agentes_ia (
                usuario_id, dispositivo_id, nombre, modelo, instrucciones, personalidad, activo, 
                descripcion_negocio, industria, objetivo, pasos_captura, skip_existing_data, 
                seguimientos, reglas_transferencia, reglas_etiquetado, config_comportamiento
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            user_id, dispositivo_id, nombre, modelo, instrucciones, personalidad, activo, 
            descripcion_negocio, industria, obj_final, pasos_captura, skip_existing_data, 
            seguimientos, reglas_transferencia, reglas_etiquetado, config_comportamiento
        ))
        
        new_id = cursor.lastrowid
        conn.commit()
        
        return jsonify({"success": True, "message": "Agente creado con éxito", "id": new_id})
    except Exception as e:
        logger.exception("Error al crear agente de IA")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>', methods=['PUT'])
@jwt_required()
def update_agente_ia(agent_id):
    role_err = require_admin_role()
    if role_err:
        return role_err
    user_id = get_jwt_identity()
    data = request.get_json() or {}
    
    dispositivo_id = data.get('dispositivo_id')
    nombre = data.get('nombre')
    modelo = data.get('modelo')
    instrucciones = data.get('instrucciones')
    personalidad = data.get('personalidad')
    activo = data.get('activo')
    descripcion_negocio = data.get('descripcion_negocio')
    industria = data.get('industria')
    objetivo = data.get('objetivo')
    
    pasos_captura = data.get('pasos_captura')
    skip_existing_data = data.get('skip_existing_data')
    seguimientos = data.get('seguimientos')
    reglas_transferencia = data.get('reglas_transferencia')
    reglas_etiquetado = data.get('reglas_etiquetado')
    config_comportamiento = data.get('config_comportamiento')
    
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Buscar el agente existente para asegurar pertenencia
        cursor.execute("SELECT * FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        existing = cursor.fetchone()
        if not existing:
            return jsonify({"success": False, "message": "Agente no encontrado o sin permisos"}), 404
            
        # Preparar campos a actualizar
        new_dispositivo = dispositivo_id if dispositivo_id is not None else existing['dispositivo_id']
        new_nombre = nombre if nombre is not None else existing['nombre']
        new_modelo = modelo if modelo is not None else existing['modelo']
        new_instrucciones = instrucciones if instrucciones is not None else existing['instrucciones']
        new_personalidad = personalidad if personalidad is not None else existing['personalidad']
        new_activo = 1 if (activo is True or activo == 1 or (activo is None and existing['activo'])) else 0
        if activo is False or activo == 0:
            new_activo = 0
        new_descripcion = descripcion_negocio if descripcion_negocio is not None else existing.get('descripcion_negocio')
        new_industria = industria if industria is not None else existing.get('industria')
        new_objetivo = objetivo if objetivo is not None else existing.get('objetivo')
        
        new_pasos = pasos_captura if pasos_captura is not None else existing.get('pasos_captura')
        new_skip = 1 if (skip_existing_data is True or skip_existing_data == 1 or (skip_existing_data is None and existing.get('skip_existing_data'))) else 0
        if skip_existing_data is False or skip_existing_data == 0:
            new_skip = 0
        new_seguimientos = seguimientos if seguimientos is not None else existing.get('seguimientos')
        new_transferencia = reglas_transferencia if reglas_transferencia is not None else existing.get('reglas_transferencia')
        new_etiquetado = reglas_etiquetado if reglas_etiquetado is not None else existing.get('reglas_etiquetado')
        new_config = config_comportamiento if config_comportamiento is not None else existing.get('config_comportamiento')

        # Preservar tokens de integraciones (Google Calendar, Calendly) que el
        # frontend no conoce ni envía: sin esto, cualquier guardado normal de
        # configuraciones sobreescribe config_comportamiento y borra los tokens
        # guardados por el callback de OAuth.
        if config_comportamiento is not None and existing.get('config_comportamiento'):
            protected_keys = [
                'google_access_token', 'google_refresh_token', 'google_token_expiry',
                'calendly_access_token', 'calendly_refresh_token', 'calendly_token_expiry',
            ]
            try:
                old_cfg = json.loads(existing['config_comportamiento'])
                new_cfg = json.loads(new_config)
                changed = False
                for key in protected_keys:
                    if key not in new_cfg and key in old_cfg:
                        new_cfg[key] = old_cfg[key]
                        changed = True
                if changed:
                    new_config = json.dumps(new_cfg)
            except Exception:
                pass

        # Validar si el plan permite objetivos de IA avanzados
        cursor.execute(
            """
            SELECT p.permite_todos_objetivos_ia
            FROM suscripciones s
            INNER JOIN planes p ON p.id = s.plan_id
            WHERE s.usuario_id = %s
            ORDER BY FIELD(s.estado, 'activa', 'prueba', 'vencida', 'cancelada'), s.fecha_vencimiento DESC, s.id DESC
            LIMIT 1
            """,
            (user_id,)
        )
        plan_row = cursor.fetchone()
        permite_todos_objetivos = bool(plan_row["permite_todos_objetivos_ia"]) if plan_row else False

        # Solo bloquear cuando el guardado intenta CAMBIAR el objetivo a uno avanzado,
        # no en cada actualización del agente (de lo contrario, un plan vencido dejaría
        # sin poder editar nada a un agente que ya tenía un objetivo avanzado asignado).
        objetivo_is_changing = objetivo is not None and objetivo != existing.get('objetivo')
        if objetivo_is_changing and not permite_todos_objetivos and new_objetivo != 'preguntas_frecuentes':
            return jsonify({"success": False, "message": "Objetivo avanzado no disponible en tu plan actual. Mejora al Plan Advanced para desbloquearlo."}), 400

        cursor.execute("""
            UPDATE agentes_ia
            SET dispositivo_id = %s, nombre = %s, modelo = %s, instrucciones = %s, personalidad = %s, activo = %s, 
                descripcion_negocio = %s, industria = %s, objetivo = %s, pasos_captura = %s, skip_existing_data = %s, 
                seguimientos = %s, reglas_transferencia = %s, reglas_etiquetado = %s, config_comportamiento = %s
            WHERE id = %s AND usuario_id = %s
        """, (
            new_dispositivo, new_nombre, new_modelo, new_instrucciones, new_personalidad, new_activo, 
            new_descripcion, new_industria, new_objetivo, new_pasos, new_skip, 
            new_seguimientos, new_transferencia, new_etiquetado, new_config, agent_id, user_id
        ))
        
        conn.commit()
        return jsonify({"success": True, "message": "Agente actualizado con éxito"})
    except Exception as e:
        logger.exception("Error al actualizar agente de IA")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>', methods=['DELETE'])
@jwt_required()
def delete_agente_ia(agent_id):
    role_err = require_admin_role()
    if role_err:
        return role_err
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Eliminar si pertenece al usuario
        cursor.execute("DELETE FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        conn.commit()
        
        if cursor.rowcount == 0:
            return jsonify({"success": False, "message": "Agente no encontrado o sin permisos"}), 404
            
        return jsonify({"success": True, "message": "Agente eliminado con éxito"})
    except Exception as e:
        logger.exception("Error al eliminar agente de IA")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/auth/google', methods=['GET'])
def auth_google():
    agent_id = request.args.get('agent_id')
    token = request.args.get('token')
    
    if not agent_id or not token:
        return "Faltan parámetros obligatorios (agent_id, token)", 400
        
    try:
        decoded = decode_token(token)
        user_id = decoded.get('sub') or decoded.get('identity')
    except Exception:
        return "Token de autorización inválido o expirado", 401
        
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
    agent = cursor.fetchone()
    cursor.close()
    conn.close()
    
    if not agent:
        return "Agente no encontrado o sin permisos", 404
        
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    redirect_uri = os.getenv("GOOGLE_REDIRECT_URI")
    
    if not client_id or not redirect_uri:
        return "Las credenciales de Google OAuth no están configuradas en el servidor", 500
        
    scope = "https://www.googleapis.com/auth/calendar.events https://www.googleapis.com/auth/userinfo.email"
    auth_url = (
        f"https://accounts.google.com/o/oauth2/v2/auth?"
        f"response_type=code&"
        f"client_id={client_id}&"
        f"redirect_uri={redirect_uri}&"
        f"scope={scope}&"
        f"access_type=offline&"
        f"prompt=consent&"
        f"state={agent_id}"
    )
    return redirect(auth_url)


@agentes_ia_blueprint.route('/api/auth/google/callback', methods=['GET'])
def auth_google_callback():
    code = request.args.get('code')
    agent_id = request.args.get('state')
    error = request.args.get('error')
    
    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:5173")
    
    if error:
        return redirect(f"{frontend_url}/agentes-ia?error={error}")
        
    if not code or not agent_id:
        return "Falta el código de autorización o el estado del agente", 400
        
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    redirect_uri = os.getenv("GOOGLE_REDIRECT_URI")
    
    try:
        token_url = "https://oauth2.googleapis.com/token"
        payload = {
            "code": code,
            "client_id": client_id,
            "client_secret": client_secret,
            "redirect_uri": redirect_uri,
            "grant_type": "authorization_code"
        }
        res = requests.post(token_url, data=payload)
        tokens = res.json()
        
        if "error" in tokens:
            return redirect(f"{frontend_url}/agentes-ia?error={tokens.get('error_description', 'Error en Google Token Exchange')}")
            
        access_token = tokens.get("access_token")
        refresh_token = tokens.get("refresh_token")
        expires_in = tokens.get("expires_in", 3600)
        
        userinfo_url = "https://www.googleapis.com/oauth2/v2/userinfo"
        userinfo_res = requests.get(userinfo_url, headers={"Authorization": f"Bearer {access_token}"})
        email = userinfo_res.json().get("email", "Cuenta de Google")
        
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT config_comportamiento FROM agentes_ia WHERE id = %s", (agent_id,))
        agent = cursor.fetchone()
        
        config = {}
        if agent and agent.get("config_comportamiento"):
            try:
                config = json.loads(agent["config_comportamiento"])
            except Exception:
                config = {}
                
        config["calGoogleConnected"] = True
        config["calGoogleEmail"] = email
        config["google_access_token"] = access_token
        if refresh_token:
            config["google_refresh_token"] = refresh_token
        config["google_token_expiry"] = time.time() + expires_in
        
        cursor.execute(
            "UPDATE agentes_ia SET config_comportamiento = %s WHERE id = %s",
            (json.dumps(config), agent_id)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return redirect(f"{frontend_url}/agentes-ia?success=true&provider=google")
    except Exception as e:
        logger.exception("Error en Google OAuth Callback")
        return redirect(f"{frontend_url}/agentes-ia?error={str(e)}")


@agentes_ia_blueprint.route('/api/auth/calendly', methods=['GET'])
def auth_calendly():
    agent_id = request.args.get('agent_id')
    token = request.args.get('token')
    
    if not agent_id or not token:
        return "Faltan parámetros obligatorios (agent_id, token)", 400
        
    try:
        decoded = decode_token(token)
        user_id = decoded.get('sub') or decoded.get('identity')
    except Exception:
        return "Token de autorización inválido o expirado", 401
        
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
    agent = cursor.fetchone()
    cursor.close()
    conn.close()
    
    if not agent:
        return "Agente no encontrado o sin permisos", 404
        
    client_id = os.getenv("CALENDLY_CLIENT_ID")
    redirect_uri = os.getenv("CALENDLY_REDIRECT_URI")
    
    if not client_id or not redirect_uri:
        return "Las credenciales de Calendly OAuth no están configuradas en el servidor", 500
        
    auth_url = (
        f"https://auth.calendly.com/oauth/authorize?"
        f"client_id={client_id}&"
        f"redirect_uri={redirect_uri}&"
        f"response_type=code&"
        f"state={agent_id}"
    )
    return redirect(auth_url)


@agentes_ia_blueprint.route('/api/auth/calendly/callback', methods=['GET'])
def auth_calendly_callback():
    code = request.args.get('code')
    agent_id = request.args.get('state')
    error = request.args.get('error')
    
    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:5173")
    
    if error:
        return redirect(f"{frontend_url}/agentes-ia?error={error}")
        
    if not code or not agent_id:
        return "Falta el código de autorización o el estado del agente", 400
        
    client_id = os.getenv("CALENDLY_CLIENT_ID")
    client_secret = os.getenv("CALENDLY_CLIENT_SECRET")
    redirect_uri = os.getenv("CALENDLY_REDIRECT_URI")
    
    try:
        token_url = "https://auth.calendly.com/oauth/token"
        payload = {
            "code": code,
            "client_id": client_id,
            "client_secret": client_secret,
            "redirect_uri": redirect_uri,
            "grant_type": "authorization_code"
        }
        res = requests.post(token_url, data=payload)
        tokens = res.json()
        
        if "error" in tokens:
            return redirect(f"{frontend_url}/agentes-ia?error={tokens.get('error_description', 'Error en Calendly Token Exchange')}")
            
        access_token = tokens.get("access_token")
        refresh_token = tokens.get("refresh_token")
        expires_in = tokens.get("expires_in", 7200)
        
        user_url = "https://api.calendly.com/users/me"
        user_res = requests.get(user_url, headers={"Authorization": f"Bearer {access_token}"})
        user_data = user_res.json()
        
        email = user_data.get("resource", {}).get("email", "Cuenta de Calendly")
        
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT config_comportamiento FROM agentes_ia WHERE id = %s", (agent_id,))
        agent = cursor.fetchone()
        
        config = {}
        if agent and agent.get("config_comportamiento"):
            try:
                config = json.loads(agent["config_comportamiento"])
            except Exception:
                config = {}
                
        config["calCalendlyConnected"] = True
        config["calCalendlyEmail"] = email
        config["calendly_access_token"] = access_token
        if refresh_token:
            config["calendly_refresh_token"] = refresh_token
        config["calendly_token_expiry"] = time.time() + expires_in
        
        cursor.execute(
            "UPDATE agentes_ia SET config_comportamiento = %s WHERE id = %s",
            (json.dumps(config), agent_id)
        )
        conn.commit()
        cursor.close()
        conn.close()
        
        return redirect(f"{frontend_url}/agentes-ia?success=true&provider=calendly")
    except Exception as e:
        logger.exception("Error en Calendly OAuth Callback")
        return redirect(f"{frontend_url}/agentes-ia?error={str(e)}")

# ==========================================
# ENDPOINTS PARA RECURSOS MULTIMEDIA DEL AGENTE
# ==========================================

from werkzeug.utils import secure_filename
import uuid
from flask import current_app

ALLOWED_RESOURCE_EXTENSIONS = {'jpg', 'jpeg', 'png', 'mp3', 'wav', 'ogg', 'mp4', 'avi', 'mov', 'pdf', 'doc', 'docx', 'txt', 'csv'}

def allowed_resource_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_RESOURCE_EXTENSIONS

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/recursos', methods=['GET'])
@jwt_required()
def get_agente_recursos(agent_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404
            
        cursor.execute("SELECT * FROM agente_recursos WHERE agente_id = %s ORDER BY creado_en DESC", (agent_id,))
        recursos = cursor.fetchall()
        
        for r in recursos:
            if r.get('creado_en'):
                r['creado_en'] = r['creado_en'].isoformat()
                
        return jsonify({"success": True, "data": recursos})
    except Exception as e:
        logger.exception("Error al listar recursos del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/recursos', methods=['POST'])
@jwt_required()
def upload_agente_recurso(agent_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404
            
        file = request.files.get('file')
        if not file or not file.filename:
            return jsonify({"success": False, "message": "Archivo requerido"}), 400
            
        tipo = request.form.get('tipo', 'Imagen')
        descripcion = request.form.get('descripcion', '')
        notas_uso = request.form.get('notas_uso', '')
        
        if not allowed_resource_file(file.filename):
            return jsonify({"success": False, "message": "Formato de archivo no permitido"}), 400
            
        # Crear carpeta de subidas de recursos
        upload_dir = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), "recursos", str(user_id))
        os.makedirs(upload_dir, exist_ok=True)
        
        filename = secure_filename(file.filename)
        unique_name = f"{uuid.uuid4().hex}_{filename}"
        file.save(os.path.join(upload_dir, unique_name))
        
        # Guardar en base de datos la URL relativa/absoluta
        media_path = f"recursos/{user_id}/{unique_name}"
        media_url = f"{request.host_url.rstrip('/')}/media/{media_path}"
        
        try:
            cursor.execute("""
                INSERT INTO agente_recursos (agente_id, tipo, archivo_url, nombre_archivo, descripcion, notas_uso)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (agent_id, tipo, media_url, filename, descripcion, notas_uso))
            conn.commit()
        except Exception as insert_err:
            if "1265" in str(insert_err) or "Data truncated" in str(insert_err) or "tipo" in str(insert_err):
                try:
                    cursor.execute("ALTER TABLE agente_recursos MODIFY COLUMN tipo VARCHAR(50) NOT NULL")
                    conn.commit()
                    cursor.execute("""
                        INSERT INTO agente_recursos (agente_id, tipo, archivo_url, nombre_archivo, descripcion, notas_uso)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    """, (agent_id, tipo, media_url, filename, descripcion, notas_uso))
                    conn.commit()
                except Exception as retry_err:
                    raise retry_err
            else:
                raise insert_err
        
        new_id = cursor.lastrowid
        
        return jsonify({
            "success": True,
            "message": "Recurso subido con éxito",
            "data": {
                "id": new_id,
                "agente_id": agent_id,
                "tipo": tipo,
                "archivo_url": media_url,
                "nombre_archivo": filename,
                "descripcion": descripcion,
                "notas_uso": notas_uso
            }
        })
    except Exception as e:
        logger.exception("Error al subir recurso del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/recursos/<int:recurso_id>', methods=['DELETE'])
@jwt_required()
def delete_agente_recurso(agent_id, recurso_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404
            
        # Obtener el recurso
        cursor.execute("SELECT * FROM agente_recursos WHERE id = %s AND agente_id = %s", (recurso_id, agent_id))
        recurso = cursor.fetchone()
        if not recurso:
            return jsonify({"success": False, "message": "Recurso no encontrado"}), 404
            
        # Eliminar archivo físico
        archivo_url = recurso.get('archivo_url', '')
        if archivo_url:
            try:
                path_part = archivo_url.split('/media/')[-1]
                file_path = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), *path_part.split('/'))
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                logger.error(f"Error al eliminar archivo físico de recurso: {e}")
                
        # Eliminar de la base de datos
        cursor.execute("DELETE FROM agente_recursos WHERE id = %s", (recurso_id,))
        conn.commit()
        
        return jsonify({"success": True, "message": "Recurso eliminado con éxito"})
    except Exception as e:
        logger.exception("Error al eliminar recurso del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# ==========================================
# ENDPOINTS PARA BASE DE CONOCIMIENTO (ENTRENAMIENTO)
# ==========================================

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/conocimiento', methods=['GET'])
@jwt_required()
def get_agente_conocimiento(agent_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404
            
        tipo = request.args.get('tipo')
        if tipo:
            cursor.execute("SELECT * FROM agente_conocimiento WHERE agente_id = %s AND tipo = %s ORDER BY creado_en DESC", (agent_id, tipo))
        else:
            cursor.execute("SELECT * FROM agente_conocimiento WHERE agente_id = %s ORDER BY creado_en DESC", (agent_id,))
            
        items = cursor.fetchall()
        for item in items:
            if item.get('creado_en'):
                item['creado_en'] = item['creado_en'].isoformat()
                
        return jsonify({"success": True, "data": items})
    except Exception as e:
        logger.exception("Error al listar conocimiento del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/conocimiento', methods=['POST'])
@jwt_required()
def add_agente_conocimiento(agent_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id, nombre FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404

        tipo = request.form.get('tipo')
        if not tipo:
            data = request.get_json() or {}
            tipo = data.get('tipo')
            titulo = data.get('titulo')
            contenido = data.get('contenido')
            url = data.get('url')
        else:
            titulo = request.form.get('titulo')
            contenido = request.form.get('contenido')
            url = request.form.get('url')
            
        if not tipo or not titulo:
            return jsonify({"success": False, "message": "Tipo y Título son requeridos"}), 400
            
        file_url = None
        if tipo == 'Doc' and 'file' in request.files:
            file = request.files['file']
            if file and file.filename:
                ALLOWED_DOC_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'csv', 'xls', 'xlsx'}
                ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
                if ext not in ALLOWED_DOC_EXTENSIONS:
                    return jsonify({"success": False, "message": "Formato de documento no permitido"}), 400
                    
                upload_dir = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), "conocimiento", str(user_id))
                os.makedirs(upload_dir, exist_ok=True)
                
                filename = secure_filename(file.filename)
                unique_name = f"{uuid.uuid4().hex}_{filename}"
                full_path = os.path.join(upload_dir, unique_name)
                file.save(full_path)
                
                media_path = f"conocimiento/{user_id}/{unique_name}"
                file_url = f"{request.host_url.rstrip('/')}/media/{media_path}"
                
                # Extraer texto de forma automática
                extracted_content = ""
                if ext == 'pdf':
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(full_path)
                        for page in reader.pages:
                            t = page.extract_text()
                            if t:
                                extracted_content += t + "\n"
                    except Exception as p_err:
                        logger.error(f"Error leyendo PDF: {p_err}")
                elif ext in ('doc', 'docx'):
                    try:
                        import docx
                        doc_obj = docx.Document(full_path)
                        for para in doc_obj.paragraphs:
                            if para.text:
                                extracted_content += para.text + "\n"
                    except Exception as d_err:
                        logger.error(f"Error leyendo DOCX: {d_err}")
                elif ext in ('txt', 'csv'):
                    try:
                        with open(full_path, "r", encoding="utf-8", errors="ignore") as f_in:
                            extracted_content = f_in.read()
                    except Exception as f_err:
                        logger.error(f"Error leyendo TXT/CSV: {f_err}")
                elif ext in ('xls', 'xlsx'):
                    try:
                        import pandas as pd
                        xl = pd.ExcelFile(full_path)
                        text_parts = []
                        for sheet_name in xl.sheet_names:
                            df = xl.parse(sheet_name)
                            df = df.dropna(how='all')
                            if not df.empty:
                                text_parts.append(f"--- Hoja: {sheet_name} ---")
                                text_parts.append(df.to_string(index=False))
                        extracted_content = "\n\n".join(text_parts)
                    except Exception as xl_err:
                        logger.error(f"Error leyendo Excel: {xl_err}")
                
                if extracted_content:
                    contenido = extracted_content

                # Si no se pudo extraer nada (ej. .doc binario viejo que python-docx no
                # sabe leer, o un PDF escaneado sin texto), no guardes el registro vacío
                # en silencio: avisa para que el usuario intente con otro archivo/formato.
                if not contenido:
                    return jsonify({
                        "success": False,
                        "message": f"No se pudo extraer texto de \"{file.filename}\". Si es un .doc antiguo, guárdalo como .docx; si es un PDF escaneado, usa uno con texto seleccionable."
                    }), 400
        elif tipo == 'Web' and url:
            url_str = str(url).strip()
            if not url_str.startswith(("http://", "https://")):
                url_str = "https://" + url_str
            scraped_text = ""
            try:
                import re
                from con_embeddings import crawl_website
                is_full_site = False
                max_pages = 1

                # Detectar si el frontend solicitó crawling de todo el sitio
                if contenido and 'Sitio completo' in str(contenido):
                    is_full_site = True
                    max_pages = 50
                    match = re.search(r'Máx:\s*(\d+)', str(contenido))
                    if match:
                        max_pages = min(int(match.group(1)), 500)

                if is_full_site:
                    logger.info(f"Iniciando crawling de sitio completo para {url_str} (Máx: {max_pages} páginas)")
                    scraped_text = crawl_website(url_str, max_pages)
                else:
                    logger.info(f"Iniciando scraping de página individual para {url_str}")
                    scraped_text = crawl_website(url_str, 1)
            except Exception as scrap_err:
                logger.error(f"Error raspando pagina web en route: {scrap_err}")

            # Si el scraping no devolvió nada, no guardes el texto del selector como si
            # fuera el contenido real: avisa al usuario en vez de entrenar al bot con basura.
            if not scraped_text:
                return jsonify({
                    "success": False,
                    "message": f"No se pudo extraer contenido de {url_str}. La página pudo bloquear el acceso, tardar demasiado, o no tener texto disponible. Intenta con otra URL o revisa que sea accesible públicamente."
                }), 400

            contenido = scraped_text
            url = url_str  # Actualizar con la URL normalizada
        elif tipo == 'Videos' and url:
            try:
                from con_embeddings import extract_youtube_video_id, get_youtube_transcript
                yt_id = extract_youtube_video_id(url)
                if yt_id:
                    lang = 'es'
                    if contenido:
                        c_lower = str(contenido).lower()
                        if 'ingl' in c_lower or 'english' in c_lower:
                            lang = 'en'
                        elif 'franc' in c_lower or 'french' in c_lower:
                            lang = 'fr'
                        elif 'alem' in c_lower or 'german' in c_lower:
                            lang = 'de'
                        elif 'portug' in c_lower or 'portuguese' in c_lower:
                            lang = 'pt'
                    logger.info(f"Obteniendo transcripción de YouTube para ID {yt_id} en idioma '{lang}'")
                    transcript = get_youtube_transcript(yt_id, lang)
                    if transcript:
                        contenido = f"--- Transcripción de Video de YouTube ({url}) ---\n{transcript}"
                    else:
                        contenido = f"Video de YouTube ({url}) sin transcripción disponible."
                else:
                    contenido = f"Enlace de YouTube no válido: {url}"
            except Exception as yt_err:
                logger.error(f"Error en endpoint transcribiendo YouTube: {yt_err}")
                    
        # Validar límite de base de conocimiento (1 MB para Starter, 10 MB para Growth)
        cursor.execute(
            """
            SELECT p.nombre AS plan_nombre
            FROM suscripciones s
            INNER JOIN planes p ON p.id = s.plan_id
            WHERE s.usuario_id = %s
            ORDER BY FIELD(s.estado, 'activa', 'prueba', 'vencida', 'cancelada'), s.fecha_vencimiento DESC, s.id DESC
            LIMIT 1
            """,
            (user_id,)
        )
        plan_row = cursor.fetchone()
        plan_nombre = plan_row["plan_nombre"] if plan_row else "Gratis"

        # Conteo del tamaño actual del conocimiento (contenido guardado para este agente)
        cursor.execute(
            """
            SELECT SUM(LENGTH(COALESCE(contenido, ''))) AS total_size 
            FROM agente_conocimiento ac
            WHERE ac.agente_id = %s
            """,
            (agent_id,)
        )
        size_row = cursor.fetchone()
        current_size = size_row["total_size"] if (size_row and size_row.get("total_size") is not None) else 0
        new_item_size = len((contenido or "").encode("utf-8"))
        total_projected_size = current_size + new_item_size

        if plan_nombre == 'Starter' and total_projected_size > 1 * 1024 * 1024:
            return jsonify({
                "success": False, 
                "message": "Has alcanzado el límite de 1 MB de base de conocimiento para este agente. Mejora tu plan para ampliar el almacenamiento."
            }), 400

        elif plan_nombre == 'Growth' and total_projected_size > 10 * 1024 * 1024:
            return jsonify({
                "success": False, 
                "message": "Has alcanzado el límite de 10 MB de base de conocimiento en tu plan actual. Mejora al Plan Advanced para ampliar el almacenamiento."
            }), 400

        cursor.execute("""
            INSERT INTO agente_conocimiento (agente_id, tipo, titulo, contenido, url)
            VALUES (%s, %s, %s, %s, %s)
        """, (agent_id, tipo, titulo, contenido, file_url or url))
        conn.commit()
        
        new_id = cursor.lastrowid
        
        try:
            from con_embeddings import sync_conocimiento_chunks
            openai_key = os.getenv("OPENAI_API_KEY")
            gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
            sync_conocimiento_chunks(cursor, conn, agent_id, new_id, contenido, url=file_url or url, gemini_key=gemini_key, openai_key=openai_key)
        except Exception as sync_err:
            logger.error(f"Error sincronizando chunks de conocimiento: {sync_err}")
            try:
                from main import enviar_notificacion_sistema
                nombre_agente = agent.get("nombre") or "tu asistente"
                enviar_notificacion_sistema(
                    int(user_id),
                    f"⚠️ El documento \"{titulo}\" se subió a *{nombre_agente}* pero no se pudo entrenar "
                    f"correctamente, así que tu asistente todavía no va a usar ese contenido para responder. "
                    f"Intenta eliminarlo y volver a subirlo; si el error persiste, escríbenos por este mismo chat."
                )
            except Exception as notif_err:
                logger.warning(f"No se pudo enviar el aviso de fallo de entrenamiento a usuario_id={user_id}: {notif_err}")
        
        return jsonify({
            "success": True,
            "message": "Entrenamiento agregado con éxito",
            "data": {
                "id": new_id,
                "agente_id": agent_id,
                "tipo": tipo,
                "titulo": titulo,
                "contenido": contenido,
                "url": file_url or url
            }
        })
    except Exception as e:
        logger.exception("Error al agregar conocimiento del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/conocimiento/<int:item_id>', methods=['DELETE'])
@jwt_required()
def delete_agente_conocimiento(agent_id, item_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404
            
        # Obtener el ítem
        cursor.execute("SELECT * FROM agente_conocimiento WHERE id = %s AND agente_id = %s", (item_id, agent_id))
        item = cursor.fetchone()
        if not item:
            return jsonify({"success": False, "message": "Entrenamiento no encontrado"}), 404
            
        url = item.get('url')
        if item.get('tipo') == 'Doc' and url and '/media/conocimiento/' in url:
            try:
                path_part = url.split('/media/')[-1]
                file_path = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), *path_part.split('/'))
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception as e:
                logger.error(f"Error al eliminar archivo físico de documento: {e}")
                
        cursor.execute("DELETE FROM agente_conocimiento WHERE id = %s", (item_id,))
        try:
            cursor.execute("DELETE FROM agente_conocimiento_chunks WHERE conocimiento_id = %s", (item_id,))
        except Exception as chunk_del_err:
            logger.error(f"Error borrando chunks al eliminar conocimiento: {chunk_del_err}")
        conn.commit()
        
        return jsonify({"success": True, "message": "Entrenamiento eliminado con éxito"})
    except Exception as e:
        logger.exception("Error al eliminar conocimiento del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

# ==========================================
# ENDPOINTS PARA INTEGRACIÓN CON GOOGLE DRIVE
# ==========================================

@agentes_ia_blueprint.route('/api/config/google', methods=['GET'])
@jwt_required()
def get_google_config():
    return jsonify({
        "success": True,
        "client_id": os.getenv("GOOGLE_CLIENT_ID"),
        "api_key": os.getenv("GOOGLE_API_KEY")
    })

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/conocimiento/google-drive', methods=['POST'])
@jwt_required()
def download_google_drive_file(agent_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Asistente no encontrado o no autorizado"}), 404
            
        data = request.get_json() or {}
        file_id = data.get('file_id')
        access_token = data.get('access_token')
        file_name = data.get('file_name', 'Archivo de Google Drive')
        
        if not file_id or not access_token:
            return jsonify({"success": False, "message": "Faltan datos requeridos (file_id o access_token)"}), 400
            
        # Descargar archivo desde la API de Google Drive
        drive_url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"
        headers = {"Authorization": f"Bearer {access_token}"}
        
        res = requests.get(drive_url, headers=headers, stream=True)
        if res.status_code != 200:
            return jsonify({"success": False, "message": f"Error al descargar desde Google Drive: {res.text}"}), 400
            
        # Guardar en local
        upload_dir = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), "conocimiento", str(user_id))
        os.makedirs(upload_dir, exist_ok=True)
        
        filename = secure_filename(file_name)
        unique_name = f"{uuid.uuid4().hex}_{filename}"
        
        file_path = os.path.join(upload_dir, unique_name)
        with open(file_path, 'wb') as f:
            for chunk in res.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
                    
        media_path = f"conocimiento/{user_id}/{unique_name}"
        file_url = f"{request.host_url.rstrip('/')}/media/{media_path}"
        
        # Extraer texto del archivo de Google Drive
        ext = file_name.rsplit('.', 1)[1].lower() if '.' in file_name else ''
        extracted_content = ""
        if ext == 'pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        extracted_content += t + "\n"
            except Exception as p_err:
                logger.error(f"Error leyendo PDF de Google Drive: {p_err}")
        elif ext in ('doc', 'docx'):
            try:
                import docx
                doc_obj = docx.Document(file_path)
                for para in doc_obj.paragraphs:
                    if para.text:
                        extracted_content += para.text + "\n"
            except Exception as d_err:
                logger.error(f"Error leyendo DOCX de Google Drive: {d_err}")
        elif ext in ('txt', 'csv'):
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f_in:
                    extracted_content = f_in.read()
            except Exception as f_err:
                logger.error(f"Error leyendo TXT/CSV de Google Drive: {f_err}")
        elif ext in ('xls', 'xlsx'):
            try:
                import pandas as pd
                xl = pd.ExcelFile(file_path)
                text_parts = []
                for sheet_name in xl.sheet_names:
                    df = xl.parse(sheet_name)
                    df = df.dropna(how='all')
                    if not df.empty:
                        text_parts.append(f"--- Hoja: {sheet_name} ---")
                        text_parts.append(df.to_string(index=False))
                extracted_content = "\n\n".join(text_parts)
            except Exception as xl_err:
                logger.error(f"Error leyendo Excel de Google Drive: {xl_err}")

        final_content = extracted_content or 'Importado desde Google Drive'
        
        # Validar límite de base de conocimiento (1 MB para Starter, 10 MB para Growth)
        cursor.execute(
            """
            SELECT p.nombre AS plan_nombre
            FROM suscripciones s
            INNER JOIN planes p ON p.id = s.plan_id
            WHERE s.usuario_id = %s
            ORDER BY FIELD(s.estado, 'activa', 'prueba', 'vencida', 'cancelada'), s.fecha_vencimiento DESC, s.id DESC
            LIMIT 1
            """,
            (user_id,)
        )
        plan_row = cursor.fetchone()
        plan_nombre = plan_row["plan_nombre"] if plan_row else "Gratis"

        cursor.execute(
            """
            SELECT SUM(LENGTH(COALESCE(contenido, ''))) AS total_size 
            FROM agente_conocimiento ac
            WHERE ac.agente_id = %s
            """,
            (agent_id,)
        )
        size_row = cursor.fetchone()
        current_size = size_row["total_size"] if (size_row and size_row.get("total_size") is not None) else 0
        new_item_size = len(final_content.encode("utf-8"))
        total_projected_size = current_size + new_item_size

        if plan_nombre == 'Starter' and total_projected_size > 1 * 1024 * 1024:
            return jsonify({
                "success": False, 
                "message": "Has alcanzado el límite de 1 MB de base de conocimiento para este agente. Mejora tu plan para ampliar el almacenamiento."
            }), 400

        elif plan_nombre == 'Growth' and total_projected_size > 10 * 1024 * 1024:
            return jsonify({
                "success": False, 
                "message": "Has alcanzado el límite de 10 MB de base de conocimiento para este agente. Mejora al Plan Advanced para ampliar el almacenamiento."
            }), 400

        # Insertar en base de datos
        cursor.execute("""
            INSERT INTO agente_conocimiento (agente_id, tipo, titulo, contenido, url)
            VALUES (%s, %s, %s, %s, %s)
        """, (agent_id, 'Doc', file_name, final_content, file_url))
        conn.commit()
        
        new_id = cursor.lastrowid
        
        try:
            from con_embeddings import sync_conocimiento_chunks
            openai_key = os.getenv("OPENAI_API_KEY")
            gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
            sync_conocimiento_chunks(cursor, conn, agent_id, new_id, final_content, url=file_url, gemini_key=gemini_key, openai_key=openai_key)
        except Exception as sync_err:
            logger.error(f"Error sincronizando chunks de Google Drive: {sync_err}")
        
        return jsonify({
            "success": True,
            "message": "Archivo de Google Drive importado con éxito",
            "data": {
                "id": new_id,
                "agente_id": agent_id,
                "tipo": "Doc",
                "titulo": file_name,
                "contenido": "Importado desde Google Drive",
                "url": file_url
            }
        })
    except Exception as e:
        logger.exception("Error al importar archivo de Google Drive")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/activity/stats', methods=['GET'])
@jwt_required()
def get_agent_activity_stats(agent_id):
    user_id = get_jwt_identity()
    period = request.args.get('period', '7dias') # 'hoy', '7dias', '30dias'
    
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar que el agente pertenece al usuario
        cursor.execute("SELECT id, dispositivo_id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Agente no encontrado"}), 404
            
        device_id = agent['dispositivo_id']
        if not device_id:
            return jsonify({"success": True, "conversations": 0, "messages_sent": 0, "pending_human": 0, "transferred": 0, "resolved": 0, "resolution_rate": 0, "timeline": []})

        # Calcular fecha de inicio
        now = datetime.now()
        if period == 'hoy':
            start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif period == '30dias':
            start_date = now - timedelta(days=30)
        else: # Default 7dias
            start_date = now - timedelta(days=7)
            
        # 1. Total conversaciones únicas interactuadas
        cursor.execute("""
            SELECT COUNT(DISTINCT chat_jid) AS total 
            FROM mensajes 
            WHERE dispositivo_id = %s AND fecha_mensaje >= %s
        """, (device_id, start_date))
        conversations = cursor.fetchone()['total'] or 0
        
        # 2. Total mensajes enviados por el bot
        cursor.execute("""
            SELECT COUNT(*) AS total 
            FROM mensajes 
            WHERE dispositivo_id = %s AND es_mio = 1 AND fecha_mensaje >= %s
        """, (device_id, start_date))
        messages_sent = cursor.fetchone()['total'] or 0
        
        # 3. Conversaciones transferidas a humanos
        # agente_asignado_id apunta a usuarios.id o agentes_ia.id, nunca a dispositivos.id;
        # hay que compararlo contra el ID de ESTE agente, no contra el dispositivo.
        cursor.execute("""
            SELECT COUNT(DISTINCT c.jid) AS total
            FROM contactos c
            INNER JOIN mensajes m ON c.jid = m.chat_jid AND c.dispositivo_id = m.dispositivo_id
            WHERE c.dispositivo_id = %s
              AND c.agente_asignado_id IS NOT NULL
              AND c.agente_asignado_id != %s
              AND m.fecha_mensaje >= %s
        """, (device_id, agent_id, start_date))
        transferred = cursor.fetchone()['total'] or 0

        # 4. Pendientes de atención humana (tienen agente asignado humano actualmente)
        cursor.execute("""
            SELECT COUNT(*) AS total
            FROM contactos
            WHERE dispositivo_id = %s
              AND agente_asignado_id IS NOT NULL
              AND agente_asignado_id != %s
        """, (device_id, agent_id))
        pending_human = cursor.fetchone()['total'] or 0
        
        # 5. Conversaciones resueltas (no transferidas a humano)
        resolved = max(0, conversations - transferred)
        
        # 6. Tasa de resolución
        resolution_rate = round((resolved / conversations * 100), 1) if conversations > 0 else 0.0
        
        # 7. Timeline de tendencia (últimos días)
        cursor.execute("""
            SELECT DATE(fecha_mensaje) AS date_label, COUNT(DISTINCT chat_jid) AS value
            FROM mensajes 
            WHERE dispositivo_id = %s AND fecha_mensaje >= %s
            GROUP BY DATE(fecha_mensaje)
            ORDER BY DATE(fecha_mensaje) ASC
        """, (device_id, start_date))
        timeline_rows = cursor.fetchall()
        
        timeline = []
        for row in timeline_rows:
            timeline.append({
                "date": row['date_label'].strftime("%m-%d") if hasattr(row['date_label'], 'strftime') else str(row['date_label']),
                "value": row['value']
            })
            
        return jsonify({
            "success": True,
            "conversations": conversations,
            "messages_sent": messages_sent,
            "pending_human": pending_human,
            "transferred": transferred,
            "resolved": resolved,
            "resolution_rate": resolution_rate,
            "timeline": timeline
        })
    except Exception as e:
        logger.exception("Error al obtener estadísticas de actividad")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/activity/conversations', methods=['GET'])
@jwt_required()
def get_agent_activity_conversations(agent_id):
    user_id = get_jwt_identity()
    filter_type = request.args.get('filter', 'Todas') # 'Todas', 'Humano', 'Lagunas'
    search_query = request.args.get('search', '').strip()
    
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar agente
        cursor.execute("SELECT id, dispositivo_id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Agente no encontrado"}), 404
            
        device_id = agent['dispositivo_id']
        if not device_id:
            return jsonify({"success": True, "data": []})

        # Construir consulta
        query = """
            SELECT DISTINCT c.id, c.jid, c.nombre, c.telefono, c.push_name, c.verified_name, c.notify_name, c.agente_asignado_id, c.ultimo_mensaje, c.actualizado_en, c.foto_perfil
            FROM contactos c
            INNER JOIN mensajes m ON c.jid = m.chat_jid AND c.dispositivo_id = m.dispositivo_id
            WHERE c.dispositivo_id = %s
        """
        params = [device_id]
        
        if filter_type == 'Humano':
            # agente_asignado_id apunta a usuarios.id o agentes_ia.id, nunca a dispositivos.id;
            # hay que compararlo contra el ID de ESTE agente, no contra el dispositivo.
            query += " AND c.agente_asignado_id IS NOT NULL AND c.agente_asignado_id != %s"
            params.append(agent_id)
        elif filter_type == 'Lagunas':
            # Consideramos lagunas: asignado al bot/vacío pero el último mensaje fue del cliente (es_mio = 0)
            query += """ AND (c.agente_asignado_id IS NULL OR c.agente_asignado_id = %s)
                         AND (SELECT es_mio FROM mensajes WHERE chat_jid = c.jid AND dispositivo_id = c.dispositivo_id ORDER BY fecha_mensaje DESC LIMIT 1) = 0"""
            params.append(agent_id)
            
        if search_query:
            query += " AND (c.nombre LIKE %s OR c.telefono LIKE %s)"
            params.append(f"%{search_query}%")
            params.append(f"%{search_query}%")
            
        query += " ORDER BY c.actualizado_en DESC LIMIT 50"
        
        cursor.execute(query, tuple(params))
        contacts = cursor.fetchall()
        
        # Formatear actualizado_en
        for c in contacts:
            if c.get('actualizado_en'):
                c['actualizado_en'] = c['actualizado_en'].isoformat()
                
        return jsonify({"success": True, "data": contacts})
    except Exception as e:
        logger.exception("Error al obtener lista de conversaciones")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()

@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/activity/conversations/<string:chat_jid>/messages', methods=['GET'])
@jwt_required()
def get_agent_conversation_messages(agent_id, chat_jid):
    user_id = get_jwt_identity()
    
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        # Verificar agente
        cursor.execute("SELECT id, dispositivo_id FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Agente no encontrado"}), 404
            
        device_id = agent['dispositivo_id']
        if not device_id:
            return jsonify({"success": True, "data": []})

        cursor.execute("""
            SELECT texto, es_mio, fecha_mensaje, tipo, url_media, mime_media, nombre_archivo 
            FROM mensajes 
            WHERE dispositivo_id = %s AND chat_jid = %s
            ORDER BY fecha_mensaje ASC LIMIT 100
        """, (device_id, chat_jid))
        messages = cursor.fetchall()
        
        for m in messages:
            if m.get('fecha_mensaje'):
                m['fecha_mensaje'] = m['fecha_mensaje'].isoformat()
                
        return jsonify({"success": True, "data": messages})
    except Exception as e:
        logger.exception("Error al obtener mensajes del chat")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@agentes_ia_blueprint.route('/api/agentes-ia/test-upload', methods=['POST'])
@jwt_required()
def test_upload_media():
    user_id = get_jwt_identity()
    file = request.files.get('file')
    if not file or not file.filename:
        return jsonify({"success": False, "message": "Archivo requerido"}), 400
        
    upload_dir = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), "tmp_test", str(user_id))
    os.makedirs(upload_dir, exist_ok=True)
    
    filename = secure_filename(file.filename)
    unique_name = f"test_{uuid.uuid4().hex}_{filename}"
    file_path = os.path.join(upload_dir, unique_name)
    file.save(file_path)
    
    mimetype = file.content_type or ""
    tipo_media = "documento"
    if mimetype.startswith('image/'):
        tipo_media = "imagen"
    elif mimetype.startswith('video/'):
        tipo_media = "video"
    elif mimetype.startswith('audio/'):
        tipo_media = "audio"
        
    relative_url = f"media/tmp_test/{user_id}/{unique_name}"
    
    return jsonify({
        "success": True,
        "url_media": relative_url,
        "nombre_archivo": filename,
        "tipo_media": tipo_media,
        "mime_media": mimetype
    })


@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/test', methods=['POST'])
@jwt_required()
def test_agent_message(agent_id):
    user_id = get_jwt_identity()
    payload = request.get_json() or {}
    message_text = payload.get('message', '').strip()
    history = payload.get('history', [])
    url_media = payload.get('url_media')
    tipo_media = payload.get('tipo_media')
    nombre_archivo = payload.get('nombre_archivo')
    mime_media = payload.get('mime_media')
    
    openai_key = os.getenv("OPENAI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    nvidia_key = os.getenv("NVIDIA_API_KEY")
    
    transcription = None
    if url_media and tipo_media == "audio":
        path_part = url_media.split('media/', 1)[1] if 'media/' in url_media else url_media
        local_path = os.path.join(current_app.config.get("UPLOAD_FOLDER", "media"), *path_part.split('/'))
        if os.path.exists(local_path):
            if gemini_key:
                from main import transcribe_audio_with_gemini
                transcription = transcribe_audio_with_gemini(local_path, gemini_key)
            if not transcription and openai_key:
                from main import transcribe_audio_with_whisper
                transcription = transcribe_audio_with_whisper(local_path, openai_key)
            
            if transcription:
                message_text = transcription
                
    if url_media and tipo_media in ("imagen", "image"):
        img_desc = f"[El cliente ha enviado una imagen: {nombre_archivo or 'Archivo'}]"
        message_text = f"{img_desc} {message_text}".strip()
    elif url_media and tipo_media == "documento":
        doc_desc = f"[El cliente ha enviado un documento: {nombre_archivo or 'Archivo'}]"
        message_text = f"{doc_desc} {message_text}".strip()
        
    if not message_text and not url_media:
        return jsonify({"success": False, "message": "Mensaje requerido"}), 400
        
    openai_key = os.getenv("OPENAI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    nvidia_key = os.getenv("NVIDIA_API_KEY")
    
    if not openai_key and not gemini_key and not nvidia_key:
        return jsonify({
            "success": True, 
            "reply": "⚠️ No hay API keys configuradas en el servidor de procesos para interactuar con el modelo de lenguaje de inteligencia artificial."
        })

    conn = None
    cursor = None
    try:
        from main import call_llm_api
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT * FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Agente no encontrado"}), 404
            
        # Base de conocimiento (RAG Vectorial con similitud coseno)
        conocimiento_text = ""
        chunks_found = []
        try:
            from con_embeddings import search_relevant_chunks
            chunks_found = search_relevant_chunks(cursor, agent["id"], message_text, top_n=5, gemini_key=gemini_key, openai_key=openai_key)
        except Exception as rag_err:
            logger.error(f"Error realizando busqueda RAG en simulador: {rag_err}")
            
        if chunks_found:
            conocimiento_text = "FRAGMENTOS DE CONOCIMIENTO RELEVANTES ENCONTRADOS:\n"
            for idx, chunk in enumerate(chunks_found):
                conocimiento_text += f"- [Fragmento {idx+1}]: {chunk}\n"
        else:
            # Fallback a volcado completo si no hay chunks
            cursor.execute("SELECT titulo, contenido, tipo FROM agente_conocimiento WHERE agente_id = %s", (agent["id"],))
            conocimiento_rows = cursor.fetchall()
            for i, item in enumerate(conocimiento_rows):
                conocimiento_text += f"\nDocumento/FAQ {i+1} ({item.get('titulo', 'Sin título')}):\n{item.get('contenido', '')}\n"

        # Cargar recursos multimedia
        cursor.execute("SELECT tipo, archivo_url, nombre_archivo, descripcion, notas_uso FROM agente_recursos WHERE agente_id = %s", (agent["id"],))
        recursos_rows = cursor.fetchall()
        recursos_text = ""
        if recursos_rows:
            recursos_text = "RECURSOS MULTIMEDIA DEL NEGOCIO:\n"
            recursos_text += "Tienes disponibles los siguientes archivos multimedia. Si el cliente te pide fotos, imágenes, audios o videos sobre algún tema listado abajo, o si consideras que enviar uno de estos archivos le ayudará a resolver su duda, debes incluir su Enlace (URL) exacto en tu respuesta de forma natural. No inventes URLs, usa solo las indicadas aquí:\n"
            for r in recursos_rows:
                desc = r.get("descripcion") or "Sin descripción"
                notas = r.get("notas_uso") or ""
                recursos_text += f"- Archivo: {r.get('nombre_archivo')} (Tipo: {r.get('tipo')}), Enlace: {r.get('archivo_url')}, Descripción: \"{desc}\""
                if notas:
                    recursos_text += f", Notas de uso: \"{notas}\""
                recursos_text += "\n"
            recursos_text += "\n"

        history_text = ""
        for h in history[-10:]:
            sender_label = "Cliente" if h.get("sender") == "user" else "Asistente"
            history_text += f"{sender_label}: {h.get('text')}\n"
        history_text += f"Cliente: {message_text}\n"

        # Pasos de captura
        pasos_captura_raw = agent.get("pasos_captura")
        pasos_text = ""
        pending_field_name = None
        if pasos_captura_raw:
            try:
                pasos = json.loads(pasos_captura_raw)
                if isinstance(pasos, list) and len(pasos) > 0:
                    pasos_text = "PASOS DE CAPTURA DE DATOS:\n"
                    pasos_text += "Debes recopilar la siguiente información del cliente de forma cálida y natural. IMPORTANTE: Si el cliente te hace una pregunta o pide información, PRIMERO responde a su duda detalladamente usando la base de conocimiento o recursos disponibles, y luego, al final de tu respuesta, solicita de forma cordial el siguiente dato de captura pendiente. Nunca ignores las preguntas del cliente para limitarte a pedir datos. Pregunta por el siguiente dato pendiente únicamente cuando el cliente responda a la pregunta anterior. Para evitar repeticiones, solicita SOLO el próximo paso pendiente, no todos los pasos a la vez:\n"

                    # Reconstruir qué variables ya se capturaron en turnos anteriores de esta
                    # misma simulación, leyendo las notas "[Simulación de Captura]" que quedaron
                    # en el historial (el simulador no tiene un contacto real en BD como en main.py).
                    captured_vars = {}
                    capture_note_re = re.compile(r'(\w+):\s*\*([^*]*)\*')
                    for h in history:
                        h_text = h.get('text') or ''
                        if 'Datos extraídos' in h_text:
                            for var_name, var_val in capture_note_re.findall(h_text):
                                if var_val.strip():
                                    captured_vars[var_name.strip().lower()] = var_val.strip()

                    pending_steps = []
                    for idx, p in enumerate(pasos):
                        var_name = (p.get("field") or "").lower().strip()
                        if var_name and var_name in captured_vars:
                            continue
                        pending_steps.append((idx, p))

                    if pending_steps:
                        next_idx, next_step = pending_steps[0]
                        pasos_text += f"- Paso {next_idx+1}: {next_step.get('text')} (Para la propiedad: {next_step.get('field')}) [PENDIENTE POR PREGUNTAR]\n"
                        pending_field_name = (next_step.get('field') or '').lower().strip()
                    else:
                        pasos_text += "- No quedan pasos de captura pendientes para este contacto.\n"
                    pasos_text += "\n"

                    pasos_text += (
                        "IMPORTANTE: NUNCA reinicies el saludo ni repitas la introducción del negocio si la "
                        "conversación ya está en curso. Antes de responder, revisa siempre el HISTORIAL DE LA "
                        "CONVERSACIÓN de abajo y continúa naturalmente desde el último mensaje, usando los datos "
                        "que el cliente ya dio anteriormente (nombre, servicio, correo, etc.) sin volver a "
                        "pedirlos ni repetir preguntas ya respondidas.\n\n"
                    )
            except Exception as pe:
                logger.error(f"Error parseando pasos_captura en simulador: {pe}")

        # Reglas de etiquetado
        reglas_etiquetado_raw = agent.get("reglas_etiquetado")
        reglas_etiquetado_text = ""
        if reglas_etiquetado_raw:
            try:
                reglas_etiquetado = json.loads(reglas_etiquetado_raw)
                if isinstance(reglas_etiquetado, list) and len(reglas_etiquetado) > 0:
                    reglas_etiquetado_text = "REGLAS DE ETIQUETADO DISPONIBLES:\n"
                    for rule in reglas_etiquetado:
                        reglas_etiquetado_text += f"- ID Regla: {rule.get('id')}, Condición del mensaje: \"{rule.get('text')}\", Etiqueta a aplicar: \"{rule.get('label') or rule.get('target')}\"\n"
                    reglas_etiquetado_text += "\n"
            except Exception as re_err:
                logger.error(f"Error parseando reglas_etiquetado en simulador: {re_err}")

        # Reglas de transferencia
        reglas_trans_raw = agent.get("reglas_transferencia")
        reglas_trans_text = ""
        if reglas_trans_raw:
            try:
                reglas_trans = json.loads(reglas_trans_raw)
                if isinstance(reglas_trans, list) and len(reglas_trans) > 0:
                    reglas_trans_text = "REGLAS DE TRANSFERENCIA/DERIVACIÓN A ASESOR HUMANO:\n"
                    for rule in reglas_trans:
                        reglas_trans_text += f"- ID Regla: {rule.get('id')}, Condición del mensaje: \"{rule.get('text')}\", Tipo: \"{rule.get('type')}\", Destino/Asesor: \"{rule.get('target')}\"\n"
                    reglas_trans_text += "\n"
            except Exception as rt_err:
                logger.error(f"Error parseando reglas_transferencia en simulador: {rt_err}")

        # Leer configuración de comportamiento y calendario
        config_raw = agent.get("config_comportamiento")
        config_json = {}
        cal_google_connected = False
        cal_calendly_connected = False
        cal_com_connected = False
        cal_google_meet = False
        cal_consultar_horarios = True
        cal_schedule_restriction = False
        cal_distribution_mode = "secuencial"
        use_emojis = True
        only_business = False
        seguimiento_enabled = False
        calendar_text = ""
        
        if config_raw:
            try:
                config_json = json.loads(config_raw)
                cal_google_connected = config_json.get("calGoogleConnected", False)
                cal_calendly_connected = config_json.get("calCalendlyConnected", False)
                cal_com_connected = bool(config_json.get("calComApiKey"))
                cal_google_meet = config_json.get("calGoogleMeet", False)
                cal_consultar_horarios = config_json.get("calConsultarHorarios", True)
                cal_schedule_restriction = config_json.get("calScheduleRestriction", False)
                cal_distribution_mode = config_json.get("calDistributionMode", "secuencial")
                use_emojis = config_json.get("useEmojis", True)
                only_business = config_json.get("onlyBusinessTopics", False)
                seguimiento_enabled = config_json.get("seguimientoInteligente", False)
                
                cal_provider = config_json.get("calProvider")
                if cal_provider:
                    cal_provider = str(cal_provider).lower()
                    if "google" in cal_provider:
                        cal_provider = "google"
                    elif "calendly" in cal_provider:
                        cal_provider = "calendly"
                    elif "cal.com" in cal_provider or "cal" in cal_provider:
                        cal_provider = "cal.com"

                # El calendario solo debe activarse mientras el objetivo del agente sea
                # "Agendar Citas", igual que en el motor de producción.
                if cal_provider and cal_provider != "ninguno" and agent.get("objetivo") == "agendar_citas":

                    cal_email = ""
                    if cal_provider == "google":
                        cal_email = config_json.get("calGoogleEmail")
                    elif cal_provider == "calendly":
                        cal_email = config_json.get("calCalendlyEmail")
                        
                    cal_asunto = config_json.get("calAsunto", "Reunion con {name}")
                    cal_reunion_desc = config_json.get("calReunionDesc", "")
                    cal_proactivas = config_json.get("calProactiveSuggestions", False)
                    cal_opciones_sugerir = config_json.get("calOptionCount", "3 opciones")
                    cal_msg_confirmacion = config_json.get("calConfirmationMsg")

                    # Log de depuración para saber si el backend está leyendo bien la conexión al calendario
                    logger.info(f"[CALENDAR_DEBUG] Agent {agent['id']} - Provider: {cal_provider}, Google Connected: {cal_google_connected}, Consultar: {cal_consultar_horarios}")

                    calendar_text = f"CALENDARIO Y RESERVAS:\n"
                    calendar_text += f"- Proveedor conectado: {cal_provider.upper()}\n"
                    if cal_email:
                        calendar_text += f"- Correo de reservas: {cal_email}\n"

                    # Antes esto usaba un nombre/correo de mentira fijos ("Cliente de
                    # Prueba"/"test@example.com") sin importar lo que la persona
                    # probando el asistente hubiera escrito, asi que la cita de prueba
                    # salia con datos genericos aunque el modelo ya supiera el nombre
                    # real por la conversacion. Ahora se le pide al propio modelo que
                    # rellene la plantilla con los datos reales que ya extrajo, en vez
                    # de que Python la rellene con datos falsos de antemano.
                    calendar_text += (
                        f"- Al crear la cita usando las herramientas de Google Calendar, arma el "
                        f"título (summary) y la descripción usando estas plantillas: título=\"{cal_asunto}\", "
                        f"descripción=\"{cal_reunion_desc}\" — reemplazando {{name}} por el NOMBRE REAL del "
                        f"cliente que hayas obtenido en esta conversación (nunca un nombre genérico como "
                        f"\"Cliente de Prueba\") y {{email}} por su correo real ya proporcionado.\n"
                    )
                    
                    if cal_proactivas:
                        calendar_text += f"- SUGERENCIA PROACTIVA: Ofrece de forma proactiva {cal_opciones_sugerir} de horarios disponibles al cliente en vez de preguntarle qué hora prefiere.\n"
                        
                    if cal_msg_confirmacion:
                        calendar_text += f"- Cuando la cita sea confirmada exitosamente, DEBES redactar tu respuesta de confirmación siguiendo exactamente esta plantilla (reemplazando los campos con los datos reales de la cita):\n{cal_msg_confirmacion}\n"
                    else:
                        calendar_text += "- Si el cliente solicita agendar una cita o reservar una hora, dile que con gusto le enviaremos la confirmación o invitación por correo electrónico.\n"
                    
                    if cal_distribution_mode:
                        calendar_text += f"- Modo de distribución de agendamientos configurado: {cal_distribution_mode.upper()}.\n"
                    calendar_text += "\n"
            except Exception as ce_err:
                logger.error(f"Error parsing config_comportamiento en simulador: {ce_err}")

        # Horarios de atención de la empresa
        working_hours_text = ""
        working_hours_raw = config_json.get("calWorkingHours")
        if working_hours_raw:
            try:
                working_hours_text = "HORARIOS DE ATENCIÓN DE LA EMPRESA:\n"
                days_map = {
                    "lunes": "Lunes",
                    "martes": "Martes",
                    "miercoles": "Miércoles",
                    "jueves": "Jueves",
                    "viernes": "Viernes",
                    "sabado": "Sábado",
                    "domingo": "Domingo"
                }
                for day_key, day_name in days_map.items():
                    day_data = working_hours_raw.get(day_key)
                    if day_data and day_data.get("active"):
                        working_hours_text += f"- {day_name}: de {day_data.get('start', '09:00')} a {day_data.get('end', '18:00')}\n"
                    else:
                        working_hours_text += f"- {day_name}: Cerrado/No disponible\n"
                working_hours_text += "\n"
            except Exception as wh_err:
                logger.error(f"Error parseando calWorkingHours en simulador: {wh_err}")

        # Servicios ofrecidos y su duracion (para que la cita se agende con la duracion
        # real de cada servicio, en vez de que la IA invente cuanto dura)
        servicios_text = ""
        servicios_duracion_map = {}
        calServicios_raw = config_json.get("calServicios")
        if calServicios_raw:
            try:
                servicios_text = "SERVICIOS OFRECIDOS Y SU DURACIÓN:\n"
                for serv in calServicios_raw:
                    nombre_serv = (serv.get("nombre") or "").strip()
                    dur_serv = serv.get("duracionMinutos")
                    if nombre_serv and dur_serv:
                        servicios_text += f"- {nombre_serv}: {dur_serv} minutos\n"
                        servicios_duracion_map[nombre_serv.strip().lower()] = int(dur_serv)
                if servicios_duracion_map:
                    servicios_text += (
                        "IMPORTANTE: cuando el cliente elija uno de estos servicios, al llamar a las "
                        "herramientas de calendario indica el nombre EXACTO del servicio en el parámetro "
                        "'servicio' (tal como aparece arriba). NUNCA calcules tú mismo cuánto dura ni la "
                        "hora de fin de la cita — el sistema usa la duración real configurada para eso.\n"
                    )
                servicios_text += "\n"
            except Exception as sv_err:
                logger.error(f"Error parseando calServicios en simulador: {sv_err}")

        comportamiento_directives = "DIRECTIVAS DE FORMATO Y COMPORTAMIENTO:\n"
        if use_emojis:
            comportamiento_directives += "- Usa emojis amigables de forma moderada en tus respuestas para ser más cercano.\n"
        else:
            comportamiento_directives += "- NO uses ningún emoji en tus respuestas bajo ninguna circunstancia.\n"
            
        if only_business:
            comportamiento_directives += (
                f"- Mantente estrictamente dentro de los temas del negocio y la base de conocimiento. "
                f"Si el cliente pregunta algo totalmente ajeno al negocio (ej. \"cómo se hace una sopa\", "
                f"\"quién es Dios\", temas de cultura general, etc.), responde EXACTAMENTE con este tipo de "
                f"mensaje (adaptándolo, pero manteniendo la idea): \"Disculpa, solo puedo ayudarte con temas "
                f"relacionados a {agent.get('nombre') or 'el negocio'}. ¿Tienes alguna consulta sobre "
                f"nuestros servicios o quieres agendar una cita?\". NUNCA respondas repitiendo la "
                f"introducción/bienvenida del negocio como si fuera tu primer mensaje — eso confunde al "
                f"cliente. Responde siempre de forma directa a lo que el cliente preguntó, aunque sea para "
                f"decirle que no puedes ayudar con ese tema.\n"
            )
        comportamiento_directives += "\n"

        # Formar instrucción de seguimiento para el prompt
        instruccion_seguimiento = ""
        if seguimiento_enabled:
            instruccion_seguimiento = (
                "5. REGLA CLAVE ENTRE AGENDAR CITAS VS SEGUIMIENTO FUTURO:\n"
                "   - Si el cliente está respondiendo qué día u hora desea para su CITA MÉDICA / SERVICIO (ej: 'Para el jueves 20 a las 10 am', 'Quiero el viernes a las 3pm', 'El 20 a las 10 am', 'A las 10 me queda bien'): ESO ES UNA ELECCIÓN DE HORARIO PARA SU CITA, NO UN SEGUIMIENTO. En ese caso, procede a reservar la cita (usando 'create_google_calendar_event' o confirmando la cita) y deja 'seguimiento.programar': false.\n"
                "   - SOLO programa un seguimiento si el cliente pide explícitamente que le escribamos, recordemos o contactemos más tarde (ej. 'escríbeme a las 1:50pm', 'háblame mañana', 'recuérdame en 2 horas', 'escríbeme después'). En ese caso indícalo en el objeto 'seguimiento' con:\n"
                "     - 'programar': true\n"
                "     - 'fecha_hora_programada': fecha y hora EXACTA que pidió el cliente en formato 'YYYY-MM-DD HH:MM:SS'.\n"
                "     - 'horas_retraso': número decimal de horas en el futuro para enviar el mensaje.\n"
                "     - 'mensaje_propuesto': el mensaje que se enviará automáticamente al cliente cuando llegue esa hora para retomar la consulta.\n"
                "     IMPORTANTE EN 'respuesta_final' cuando pide recordatorio: respóndele confirmando cordialmente que le escribirás a la hora solicitada.\n"
                "   Si no solicita contacto futuro o solo está eligiendo horario de cita, deja 'programar' en false.\n"
            )
        else:
            instruccion_seguimiento = "5. Deja el objeto 'seguimiento' con 'programar': false y los demás campos en null.\n"

        tool_instructions = ""
        if agent.get("objetivo") != "agendar_citas":
            pass
        elif not cal_consultar_horarios:
            tool_instructions = (
                "IMPORTANTE - CONSULTA DE HORARIOS DESHABILITADA:\n"
                "La consulta de disponibilidad en tiempo real está deshabilitada en la configuración del asistente.\n"
                "NO utilices ninguna herramienta para consultar disponibilidad ni agendar eventos.\n"
                "Si el cliente solicita agendar una cita o ver horarios disponibles, ofrécele directamente el enlace de reserva de tu proveedor de calendario o pídele que te indique qué día y hora prefiere para reservarlo manualmente.\n\n"
            )
        else:
            if cal_provider == "google" and cal_google_connected:
                tool_instructions = (
                    "HERRAMIENTAS DE GOOGLE CALENDAR:\n"
                    "Tienes acceso a las siguientes herramientas para consultar disponibilidad y agendar citas:\n"
                    "- 'list_google_calendar_slots': Sirve para consultar la disponibilidad real. Parámetros: start_time (ISO 8601 string, UTC), end_time (ISO 8601 string, UTC), servicio (string, opcional — el nombre EXACTO del servicio que quiere el cliente, tal como aparece en SERVICIOS OFRECIDOS; así el sistema solo te devuelve huecos que realmente alcancen para ese servicio).\n"
                    "- 'create_google_calendar_event': Sirve para reservar una cita NUEVA (el cliente no tiene ninguna cita previa, o quiere agendar una adicional). Parámetros: summary (string), start_time (ISO 8601 string, UTC — la hora de INICIO que eligió el cliente), servicio (string — el nombre EXACTO del servicio elegido, tal como aparece en SERVICIOS OFRECIDOS; el sistema calcula la hora de fin automáticamente según su duración real, TÚ NO la calcules ni la envíes), attendee_email (string, opcional), description (string). Si no hay ninguna lista de SERVICIOS OFRECIDOS configurada, en su lugar envía tú el parámetro end_time (ISO 8601, UTC) calculando una duración razonable de 1 hora.\n"
                    "- 'reschedule_google_calendar_event': Sirve para CAMBIAR la fecha/hora de una cita que el cliente YA TIENE agendada (por ejemplo: \"¿me puedes cambiar mi cita para otro día?\", \"quiero mover mi cita a las 10am\"). Parámetros: start_time (ISO 8601 string, UTC — la NUEVA hora deseada), servicio (string, opcional pero recomendado — el mismo servicio de la cita original, para calcular la duración correcta), end_time (opcional, solo si no hay servicios configurados). IMPORTANTE: NUNCA uses 'create_google_calendar_event' para esto — si lo haces, se crea una cita DUPLICADA y la vieja queda sin cancelar. Usa siempre 'reschedule_google_calendar_event' cuando el cliente ya tenía una cita y solo quiere cambiar el horario.\n"
                    "- 'cancel_google_calendar_event': Sirve para CANCELAR/ELIMINAR una cita que el cliente YA TIENE agendada (por ejemplo: \"quiero cancelar mi cita\", \"ya no voy a poder ir\"). No recibe parámetros. NUNCA le digas al cliente que su cita fue cancelada sin haber llamado antes a esta herramienta y haber recibido un resultado exitoso — decir que se canceló sin ejecutar esta herramienta deja la cita real intacta en el calendario, generando confusión y una cita a la que el cliente no asistirá.\n\n"

                    "Si el cliente desea agendar una cita o saber si hay disponibilidad, DEBES ejecutar la herramienta correspondiente (incluso si en el historial de la conversación le dijiste al cliente que había un inconveniente técnico o error, debes ignorar eso e intentar llamar a la herramienta de todas formas) respondiendo ÚNICAMENTE con un JSON que contenga la propiedad 'tool_call':\n"
                    "{\n"
                    "  \"tool_call\": {\n"
                    "    \"name\": \"nombre_de_la_herramienta\",\n"
                    "    \"arguments\": { ... }\n"
                    "  }\n"
                    "}\n"
                    "IMPORTANTE: Si el cliente pregunta en general por disponibilidad para 'esta semana', 'estos días' o de forma abierta, calcula el rango usando la 'Fecha y hora actual del negocio' y llama a 'list_google_calendar_slots' abarcando los siguientes 7 días automáticamente (start_time = ahora, end_time = ahora + 7 días). No le preguntes primero al cliente el día exacto si puedes consultar tú mismo la semana completa para ofrecerle opciones. Cuando ejecutes una herramienta, NO respondas nada más (deja el resto de campos como respuesta_final en blanco o null). Solo cuando tengas los resultados de la herramienta en tu contexto, podrás responder al cliente.\n"
                    "PRIORIDAD: si hay PASOS DE CAPTURA pendientes (nombre, servicio o correo aún no obtenidos), esos van SIEMPRE primero — aunque el primer mensaje del cliente mencione la palabra 'horarios' o 'disponibilidad' (ej. el mensaje de bienvenida de un enlace/Walink), NO consultes el calendario todavía. Primero pide el dato pendiente; solo cuando ya tengas nombre, servicio Y correo, ahí sí puedes consultar o agendar en el calendario.\n"
                    "IMPORTANTE SOBRE EL LENGUAJE: si el cliente propuso un día/hora pero TODAVÍA faltan datos suyos (nombre, servicio o correo) antes de poder agendar de verdad, NUNCA uses palabras como 'confirmamos', 'confirmada' o 'tu cita ha sido agendada' — eso sugiere que ya quedó lista cuando en realidad falta un paso. En su lugar di algo como 'Para agendar tu cita del [día] a las [hora] todavía necesito tu nombre y correo electrónico'. Reserva 'confirmada'/'agendada' únicamente para cuando ya ejecutaste con éxito la herramienta de calendario.\n"
                    "PROHIBIDO INVENTAR EL CORREO: el parámetro 'attendee_email' (o 'datos_extraidos.correo') DEBE ser exactamente el texto que el cliente escribió — NUNCA lo deduzcas ni lo armes a partir de su nombre (ej. NUNCA generar 'nombre.apellido@gmail.com' solo porque el cliente te dio su nombre). Si el cliente solo dio su nombre y el correo sigue sin aparecer literalmente en sus mensajes, el correo sigue PENDIENTE — vuelve a pedírselo explícitamente, no lo agendes ni lo inventes.\n\n"

                    "IMPORTANTE SOBRE CÓMO USAR EL RESULTADO DE 'list_google_calendar_slots': el resultado trae 'disponibilidad_por_dia', una lista con UN elemento por cada día del rango consultado, cada uno con 'fecha', 'dia_semana', 'abierto' (true/false) y 'franjas_libres' (lista de horas realmente libres en formato \"HH:MM-HH:MM\", ya calculadas restando el horario de atención, las citas ya ocupadas Y la hora actual si el día es hoy — NUNCA las recalcules ni inventes otras). NUNCA ofrezcas 2 o 3 horarios del MISMO día como si fueran opciones distintas — eso confunde al cliente. En vez de eso, recorre 'disponibilidad_por_dia' completo (incluyendo el día de HOY si aparece con 'abierto': true y franjas libres, no lo saltes) y ofrécele al cliente un panorama real de VARIOS días distintos con al menos una franja libre. Si un día tiene 'abierto': false o 'franjas_libres' vacío, NO lo ofrezcas.\n"
                    "FORMATO DE LA RESPUESTA: nunca amontones las opciones en una sola oración larga con comas. Preséntalas como una listita ordenada, un renglón por día, así (usa saltos de línea reales):\n"
                    "📅 Estos son los horarios disponibles:\n"
                    "- Martes 11: 11:00 a 18:00\n"
                    "- Miércoles 12: 9:00 a 15:00\n"
                    "- Jueves 13: 9:00 a 18:00\n"
                    "¿Cuál día y hora te viene mejor?\n"
                    "Si un día tiene varias franjas libres separadas (por ejemplo por una cita en medio del día), muéstralas todas en el mismo renglón de ese día separadas por coma.\n"
                    "Si TODOS los días del rango están sin franjas libres, dile al cliente honestamente que no hay disponibilidad en esas fechas y ofrécele consultar la semana siguiente.\n"
                    "SI EL CLIENTE PIDE UNA HORA ESPECÍFICA QUE NO ALCANZA: cuando el cliente pida un horario puntual (ej. \"a las 9\") y esa hora no aparezca dentro de ninguna franja libre, NUNCA le digas solo \"no está disponible\" sin explicar por qué — eso lo deja confundido. Explícale brevemente el motivo real usando la duración de su servicio (de SERVICIOS OFRECIDOS) y por qué no alcanza (por ejemplo, porque el hueco libre antes de otra cita es más corto que la duración del servicio), y luego ofrécele las franjas reales que sí tiene disponibles ese día. Ejemplo: \"A las 9:00 no alcanza porque tu Blanqueamiento dura 1h30 y a las 10:00 ya tienes otra cita agendada — pero desde las 11:30 tienes espacio de sobra.\"\n\n"
                )

            elif cal_provider == "calendly" and cal_calendly_connected:
                tool_instructions = (
                    "HERRAMIENTAS DE CALENDLY:\n"
                    "Tienes acceso a la siguiente herramienta para obtener los enlaces de reserva y tipos de cita configurados:\n"
                    "- 'list_calendly_slots': Sirve para listar las opciones de citas (reuniones) y sus enlaces URL correspondientes para que se los envíes al cliente para que reserve directamente. No requiere parámetros.\n\n"
                    "Si el cliente desea agendar o solicita un enlace para reservar, DEBES ejecutar la herramienta correspondiente respondiendo ÚNICAMENTE con un JSON que contenga la propiedad 'tool_call':\n"
                    "{\n"
                    "  \"tool_call\": {\n"
                    "    \"name\": \"list_calendly_slots\",\n"
                    "    \"arguments\": {}\n"
                    "  }\n"
                    "}\n"
                    "IMPORTANTE: Cuando ejecutes una herramienta, NO respondas nada más. Solo cuando tengas los resultados de la herramienta con los enlaces, podrás responder al cliente enviándole el enlace del tipo de cita correspondiente.\n\n"
                )
            elif cal_provider == "cal.com" and cal_com_connected:
                tool_instructions = (
                    "HERRAMIENTAS DE CAL.COM:\n"
                    "Tienes acceso a la siguiente herramienta para obtener los enlaces de reserva y tipos de cita configurados:\n"
                    "- 'list_calcom_slots': Sirve para listar las opciones de citas (reuniones) y sus enlaces URL correspondientes de Cal.com para que se los envíes al cliente para que reserve directamente. No requiere parámetros.\n\n"
                    "Si el cliente desea agendar o solicita un enlace para reservar, DEBES ejecutar la herramienta correspondiente respondiendo ÚNICAMENTE con un JSON que contenga la propiedad 'tool_call':\n"
                    "{\n"
                    "  \"tool_call\": {\n"
                    "    \"name\": \"list_calcom_slots\",\n"
                    "    \"arguments\": {}\n"
                    "  }\n"
                    "}\n"
                    "IMPORTANTE: Cuando ejecutes una herramienta, NO respondas nada más. Solo cuando tengas los resultados de la herramienta con los enlaces, podrás responder al cliente enviándole el enlace del tipo de cita correspondiente.\n\n"
                )

        if cal_schedule_restriction and agent.get("objetivo") == "agendar_citas":
            tool_instructions += (
                "RESTRICCIÓN DE HORARIOS:\n"
                "- Al proponer o sugerir horarios libres al cliente, hazlo estrictamente en horas punto o enteras (ejemplo: 09:00, 10:00, 15:00) y evita ofrecer minutos fraccionados (como 09:30 o 14:15).\n\n"
            )

        # Obtener fecha y hora local del negocio según la zona horaria
        _DIAS_SEMANA_ES = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]
        selected_timezone = config_json.get("selectedTimezone") or "America/Guayaquil"
        local_time_str = ""
        local_dt_obj = None
        if selected_timezone:
            try:
                import pytz
                from datetime import datetime
                tz = pytz.timezone(selected_timezone)
                local_dt_obj = datetime.now(tz)
                local_time_str = local_dt_obj.strftime("%Y-%m-%d %H:%M:%S (%Z)")
            except Exception as tz_err:
                logger.error(f"Error calculando hora local del negocio: {tz_err}")

        if not local_time_str:
            from datetime import datetime
            local_dt_obj = datetime.now()
            local_time_str = local_dt_obj.strftime("%Y-%m-%d %H:%M:%S (Local Server)")

        fechas_referencia_str = ""
        if local_dt_obj is not None:
            local_time_str = f"{_DIAS_SEMANA_ES[local_dt_obj.weekday()]}, {local_time_str}"
            from datetime import timedelta
            hoy_date = local_dt_obj.date()
            _lineas_ref = []
            for idx, nombre_dia in enumerate(_DIAS_SEMANA_ES):
                days_ahead = (idx - hoy_date.weekday()) % 7
                fecha_calc = hoy_date + timedelta(days=days_ahead)
                etiqueta = " (hoy)" if days_ahead == 0 else (" (mañana)" if days_ahead == 1 else "")
                _lineas_ref.append(f"  - {nombre_dia}: {fecha_calc.isoformat()}{etiqueta}")
            fechas_referencia_str = "\n".join(_lineas_ref)

        system_prompt = (
            f"Eres el agente de IA '{agent.get('nombre') or 'Asistente'}' para WhatsApp de un negocio.\n"
            f"Fecha y hora actual del negocio: {local_time_str}\n"
            f"IMPORTANTE SOBRE FECHAS: cuando el cliente mencione un dia relativo (\"el viernes\", \"manana\", \"la proxima semana\"), usa DIRECTAMENTE la fecha exacta de la tabla de abajo — ya esta calculada, NO la vuelvas a calcular ni la cuentes tu mismo.\n"
            f"REFERENCIA DE FECHAS (ya calculadas, solo copia la que corresponda):\n{fechas_referencia_str}\n"
            f"IMPORTANTE SOBRE LA HORA AL CONFIRMAR: al escribirle al cliente (en 'respuesta_final', 'datos_extraidos' o el mensaje de confirmacion), usa SIEMPRE la hora local en formato de 12 horas que el cliente pidio (ej. \"3:00 PM\"), NUNCA la conviertas a UTC ni le muestres el desfase de zona horaria (-05:00 u otro). La hora con offset UTC es solo para el parametro tecnico start_time/end_time de la herramienta del calendario — jamas para lo que lee el cliente.\n"
            f"Tu industria/giro del negocio es: {agent.get('industria') or 'Servicios'}.\n"
            f"Tu objetivo principal es: {agent.get('objetivo') or 'Ayudar al cliente'}.\n\n"
            f"INSTRUCCIONES DE COMPORTAMIENTO:\n"
            f"{agent.get('instrucciones') or 'Responde cordialmente'}\n\n"
            + (
                f"INFORMACIÓN COMPLETA DEL NEGOCIO (usa esta información para responder con exactitud):\n"
                f"{agent.get('descripcion_negocio')}\n\n"
                if agent.get('descripcion_negocio') else ""
            ) +
            f"Tu Personalidad:\n{agent.get('personalidad', 'Amigable, profesional y servicial')}\n\n"
            f"{pasos_text}"
            f"{reglas_etiquetado_text}"
            f"{reglas_trans_text}"
            f"{calendar_text}"
            f"{working_hours_text}"
            f"{servicios_text}"
            f"{comportamiento_directives}"
            f"{recursos_text}"
            f"{tool_instructions}"
            f"BASE DE CONOCIMIENTO:\n"
            f"{conocimiento_text}\n\n"
            f"HISTORIAL DE LA CONVERSACIÓN:\n"
            f"{history_text}\n"
            "INSTRUCCIÓN DE RETORNO (DEBES CUMPLIR ESTO ESTRICTAMENTE):\n"
            "Tu tarea es analizar el último mensaje del cliente y determinar:\n"
            "1. Si coincide con alguna condición de las REGLAS DE ETIQUETADO DISPONIBLES. Si es así, incluye una lista de los IDs de las reglas que se cumplieron en 'tags_a_aplicar_ids'. Si ninguna coincide, deja un arreglo vacío [].\n"
            "2. Si coincide con alguna condición de las REGLAS DE TRANSFERENCIA A ASESOR HUMANO. Si es así, incluye el ID de la primera regla que se cumpla en 'regla_transferencia_id' (como número). Si ninguna coincide, deja este campo en null.\n"
            "3. Si el usuario proporciona datos para alguna variable pendiente en PASOS DE CAPTURA DE DATOS. Si es así, extráelos en el objeto 'datos_extraidos' con la estructura {variable: valor}. Si no hay datos nuevos, deja un objeto vacío {}.\n"
            "4. Generar la respuesta final amigable y profesional para el cliente, redactada en español, y colocarla en el campo 'respuesta_final'.\n"
            "5. Si la situación del cliente amerita enviar un archivo multimedia (imagen, audio, video, documento) de los indicados en RECURSOS MULTIMEDIA DEL NEGOCIO, coloca el valor exacto del campo 'Enlace' de ese recurso en 'url_media_a_enviar' y su tipo en 'tipo_media_a_enviar' (imagen/audio/video/documento). IMPORTANTE: cuando pongas una URL en 'url_media_a_enviar', NO la repitas ni la menciones en 'respuesta_final'. El sistema enviará el archivo automáticamente; en 'respuesta_final' escribe solo el texto conversacional natural, por ejemplo '¡Aquí tienes nuestro catálogo! 😊 ¿Te puedo ayudar con algo más?'. Si no corresponde enviar ningún recurso, deja 'url_media_a_enviar' y 'tipo_media_a_enviar' en null.\n"
            f"{instruccion_seguimiento}\n\n"
            "DEBES RESPONDER EXCLUSIVAMENTE CON UN OBJETO JSON VÁLIDO. No agregues texto antes ni después del bloque JSON, ni uses bloques de código markdown como ```json.\n"
            "Estructura del JSON esperada:\n"
            "{\n"
            "  \"tags_a_aplicar_ids\": [],\n"
            "  \"regla_transferencia_id\": null,\n"
            "  \"datos_extraidos\": {},\n"
            "  \"respuesta_final\": \"Texto de la respuesta para el cliente (puedes dejarlo vacío/null si vas a usar tool_call)\",\n"
            "  \"url_media_a_enviar\": null,\n"
            "  \"tipo_media_a_enviar\": null,\n"
            "  \"tool_call\": null,\n"
            "  \"seguimiento\": {\n"
            "     \"programar\": false,\n"
            "     \"fecha_hora_programada\": null,\n"
            "     \"horas_retraso\": null,\n"
            "     \"mensaje_propuesto\": null\n"
            "  }\n"
            "}\n"
            "IMPORTANTE: Si necesitas ejecutar una herramienta (ej. list_google_calendar_slots para ver disponibilidad, o create_google_calendar_event para agendar), rellena la propiedad 'tool_call' con la siguiente estructura:\n"
            "{\n"
            "  ... (demás campos en vacio/null) ...\n"
            "  \"tool_call\": {\n"
            "     \"name\": \"list_google_calendar_slots o create_google_calendar_event\",\n"
            "     \"arguments\": {\n"
            "        \"start_time\": \"ISO 8601 string\",\n"
            "        \"end_time\": \"ISO 8601 string\"\n"
            "     }\n"
            "  }\n"
            "}"
        )

        # Loop de ejecución del Agente en el Simulador (ReAct)
        tool_results_context = ""
        response_text = ""
        parsed_ok = False
        res_data = {}
        notes = []
        errors_list = []
        retry_nudge = ""
        json_retry_used = False
        tool_call_retry_used = False
        booking_retry_used = False
        today_omitted_retry_used = False

        for iteration in range(6):
            current_prompt = system_prompt
            if tool_results_context:
                current_prompt += f"\n\nRESULTADOS DE HERRAMIENTAS EJECUTADAS:\n{tool_results_context}\nUsa esta información para responder al cliente de manera precisa."
            if retry_nudge:
                current_prompt += retry_nudge
                retry_nudge = ""

            response_text, errors_list = call_llm_api(
                current_prompt,
                f"Prueba Asistente - {agent.get('nombre')}",
                openai_key,
                gemini_key,
                nvidia_key,
                model_override=agent.get('modelo'),
                return_errors=True
            )

            if not response_text:
                break

            # Log de depuración para ver la respuesta exacta de la IA
            logger.info(f"[LLM_RESPONSE] raw: {response_text}")

            json_match = re.search(r'\{.*\}', response_text.strip(), re.DOTALL)
            if json_match:
                try:
                    res_data = json.loads(json_match.group(0))
                    parsed_ok = True
                except Exception as pe:
                    logger.error(f"Error parseando respuesta del simulador en iteracion {iteration}: {pe}. Respuesta: {response_text}")
                    res_data = {"respuesta_final": response_text.strip()}
                    parsed_ok = True
                    break
            elif not json_retry_used and iteration < 2:
                # El modelo respondio en texto de conversacion en vez del JSON exigido
                # (le pasa mas seguido al modelo propio/Qwen que a Gemini, por su modo de
                # "pensar en voz alta"). Antes esto se mostraba tal cual al cliente y el
                # flujo de agendar citas se quedaba pegado en un mensaje de relleno tipo
                # "un momento, por favor" sin llegar a consultar el calendario. Le damos
                # una sola oportunidad de corregirse antes de rendirnos.
                json_retry_used = True
                retry_nudge = (
                    "\n\nRECORDATORIO CRITICO: tu respuesta anterior no fue el JSON exigido. "
                    "Responde EXCLUSIVAMENTE con el objeto JSON solicitado, sin ningun texto "
                    "de conversacion antes ni despues."
                )
                continue
            else:
                res_data = {"respuesta_final": response_text.strip()}
                parsed_ok = True

            # Red de seguridad: si el cliente pregunta explicitamente por horarios/
            # disponibilidad pero la IA respondio sin llamar a la herramienta de
            # calendario (le pasa a veces al modelo propio, que se salta el paso y
            # solo pregunta "que dia prefieres" sin consultar nada real), le damos
            # una sola oportunidad de corregirse antes de aceptar su respuesta.
            if (
                parsed_ok and not res_data.get("tool_call")
                and agent.get("objetivo") == "agendar_citas" and cal_consultar_horarios
                and not pending_field_name
                and not tool_call_retry_used and iteration < 2 and not tool_results_context
                and re.search(r'horari|disponib', message_text or '', re.IGNORECASE)
            ):
                tool_call_retry_used = True
                retry_nudge = (
                    "\n\nRECORDATORIO CRITICO: el cliente esta preguntando por horarios o "
                    "disponibilidad. NO le preguntes que dia o que hora prefiere sin antes "
                    "haber consultado el calendario real. DEBES llamar ahora a la herramienta "
                    "'list_google_calendar_slots' respondiendo unicamente con el JSON de 'tool_call'."
                )
                continue

            # Red de seguridad: si el cliente propuso una fecha/hora concreta para
            # agendar (ej. "martes 18 a las 4 pm") y ya se le pidieron nombre/servicio/
            # correo, pero la IA respondio SIN llamar a create_google_calendar_event
            # (a veces solo repite de memoria un rechazo anterior — "ese horario no
            # esta disponible" — sin volver a consultar el calendario real cada vez),
            # se le obliga a intentarlo de verdad antes de aceptar su respuesta.
            if (
                parsed_ok and not res_data.get("tool_call")
                and agent.get("objetivo") == "agendar_citas" and cal_consultar_horarios
                and not pending_field_name
                and not booking_retry_used and iteration < 3 and not tool_results_context
                and re.search(r'\d{1,2}\s*(:\d{2})?\s*(am|pm)|a las \d{1,2}', message_text or '', re.IGNORECASE)
            ):
                booking_retry_used = True
                retry_nudge = (
                    "\n\nRECORDATORIO CRITICO: el cliente propuso una fecha y hora concreta "
                    "para agendar. NUNCA repitas de memoria un rechazo que hayas dado antes en "
                    "esta conversacion — cada vez que el cliente proponga un horario, DEBES "
                    "llamar a la herramienta 'create_google_calendar_event' para verificar la "
                    "disponibilidad real en ese momento, respondiendo unicamente con el JSON de "
                    "'tool_call'. Si ya lo intentaste antes y fallo, intentalo de nuevo ahora, no "
                    "asumas que seguira fallando."
                )
                continue

            # Red de seguridad: se confirmo en logs que la IA a veces calcula bien la
            # disponibilidad de HOY (Python la incluye correctamente en el resultado de
            # la herramienta) pero al redactar la lista para el cliente la omite, aunque
            # la instruccion ya le dice explicitamente que no lo haga. Si el resultado de
            # la herramienta ya ejecutada incluye el dia de hoy con franjas libres, pero
            # la fecha de hoy no aparece en la respuesta final, se le pide corregirlo.
            if (
                parsed_ok and not res_data.get("tool_call")
                and agent.get("objetivo") == "agendar_citas"
                and not today_omitted_retry_used and iteration < 5 and tool_results_context
                and re.search(r'"fecha":\s*"' + re.escape(hoy_date.isoformat()) + r'",\s*"dia_semana":\s*"\w+",\s*"abierto":\s*true,\s*"franjas_libres":\s*\[(?!\])', tool_results_context)
                and hoy_date.isoformat() not in (res_data.get("respuesta_final") or "")
                and str(hoy_date.day) not in (res_data.get("respuesta_final") or "")
            ):
                today_omitted_retry_used = True
                retry_nudge = (
                    f"\n\nRECORDATORIO CRITICO: en el resultado de la herramienta que ya ejecutaste, el día de HOY "
                    f"({hoy_date.isoformat()}) aparece con 'abierto': true y franjas libres — pero tu respuesta anterior "
                    f"lo omitió de la lista de horarios disponibles. Vuelve a redactar 'respuesta_final' incluyendo TAMBIÉN "
                    f"el día de hoy con su franja libre real, junto con los demás días. No repitas la llamada a la "
                    f"herramienta, solo corrige el texto de la respuesta usando los datos que ya tienes."
                )
                continue

            # Si no hay llamada a herramienta, es la respuesta final, salimos del loop.
            # Ignora también cualquier tool_call si el objetivo ya no es "Agendar Citas".
            if not parsed_ok or "tool_call" not in res_data or not res_data["tool_call"] or agent.get("objetivo") != "agendar_citas":
                break

            tool_call = res_data["tool_call"]
            tool_name = tool_call.get("name")
            tool_args = tool_call.get("arguments") or {}

            tool_result = ""

            def _resolver_servicio_pedido(args):
                # La IA a veces se olvida de mandar el parametro "servicio" al llamar
                # a las herramientas de calendario (le pasa seguido al modelo propio),
                # y sin ese dato el filtro de duracion no funciona (puede ofrecer o
                # agendar en un hueco mas corto de lo que el servicio realmente dura).
                # No confiar solo en que la IA lo mande: se recupera del historial de
                # la conversacion si hace falta, igual que se hace con el nombre/email.
                directo = (args.get("servicio") or "").strip().lower()
                if directo:
                    return directo
                for _h in reversed(history):
                    _h_text = _h.get('text') or ''
                    if 'servicio' in _h_text.lower() and 'Datos extraídos' in _h_text:
                        _m = re.search(r'[Ss]ervicio:\s*\*([^*]+)\*', _h_text)
                        if _m and _m.group(1).strip():
                            return _m.group(1).strip().lower()
                return ""

            def _lookup_duracion_servicio(servicio_pedido):
                # No confiar en que la IA repita el nombre EXACTO del servicio configurado
                # (ej. "blanqueamiento dental" cuando el servicio se llama solo
                # "Blanqueamiento") — sin esta tolerancia la busqueda exacta fallaba y la
                # cita nunca se creaba. Se intenta primero un match exacto y, si no hay,
                # se acepta que uno sea substring del otro.
                if not servicio_pedido:
                    return None
                if servicio_pedido in servicios_duracion_map:
                    return servicios_duracion_map[servicio_pedido]
                for nombre_cfg, dur in servicios_duracion_map.items():
                    if nombre_cfg in servicio_pedido or servicio_pedido in nombre_cfg:
                        return dur
                return None

            if cal_provider == "google" and cal_google_connected:
                try:
                    from calendar_tools import refresh_google_oauth_token, list_google_calendar_events, create_google_calendar_event
                    active_token = refresh_google_oauth_token(cursor, conn, agent["id"], config_json)
                    if active_token:
                        if tool_name == "list_google_calendar_slots":
                            start_time = tool_args.get("start_time")
                            end_time = tool_args.get("end_time")
                            servicio_pedido = _resolver_servicio_pedido(tool_args)
                            duracion_pedida = _lookup_duracion_servicio(servicio_pedido)
                            if start_time and end_time:
                                res_slots = list_google_calendar_events(active_token, start_time, end_time)
                                if res_slots.get("success"):
                                    # No le pedimos a la IA que calcule la disponibilidad real
                                    # (eso causaba que ofreciera el mismo dia repetido 3 veces).
                                    # Se calcula aqui, en Python, dia por dia, cruzando el horario
                                    # de atencion configurado con los eventos ya ocupados.
                                    from calendar_tools import compute_free_slots_by_day
                                    from datetime import datetime as _dt, timedelta as _td
                                    try:
                                        _range_start = _dt.fromisoformat(start_time.replace("Z", "+00:00")).date()
                                        _range_end = _dt.fromisoformat(end_time.replace("Z", "+00:00")).date()
                                        _num_days = (_range_end - _range_start).days + 1
                                    except Exception:
                                        _range_start = hoy_date
                                        _num_days = 7
                                    disponibilidad = compute_free_slots_by_day(
                                        res_slots.get("busy_slots"),
                                        working_hours_raw,
                                        selected_timezone,
                                        _range_start,
                                        _num_days,
                                        now_dt=local_dt_obj,
                                        min_duration_minutes=duracion_pedida
                                    )
                                    tool_result = json.dumps({
                                        "success": True,
                                        "disponibilidad_por_dia": disponibilidad
                                    })
                                else:
                                    tool_result = json.dumps(res_slots)
                            else:
                                tool_result = '{"error": "Faltan parametros start_time o end_time"}'
                        elif tool_name == "create_google_calendar_event":
                            summary = tool_args.get("summary") or "Cita de Simulacion"
                            start_time = tool_args.get("start_time")
                            end_time = tool_args.get("end_time")
                            attendee_email = tool_args.get("attendee_email")
                            # Red de seguridad CRITICA: no confiar en que el correo que la IA paso
                            # en el tool_call sea real — a veces INVENTA un correo con pinta valida
                            # (nombre.apellido@gmail.com) sin que el cliente lo haya escrito nunca.
                            # Solo se acepta si aparece literalmente en el historial de la conversacion.
                            if attendee_email:
                                _historial_texto_completo = " ".join([(h.get("text") or "") for h in (history or [])]).lower()
                                if attendee_email.strip().lower() not in _historial_texto_completo:
                                    logger.warning(f"Agente {agent.get('nombre')} (simulador): se descarto un correo inventado por la IA en el tool_call ('{attendee_email}'); no aparece en el historial real del cliente.")
                                    attendee_email = None
                            description = tool_args.get("description") or "Cita agendada en simulacion"
                            servicio_pedido = _resolver_servicio_pedido(tool_args)
                            duracion_pedida = _lookup_duracion_servicio(servicio_pedido)
                            if start_time and not end_time and duracion_pedida:
                                # No confiar en que la IA calcule la hora de fin: se calcula
                                # aqui en Python usando la duracion real configurada para el
                                # servicio elegido.
                                try:
                                    from datetime import datetime as _dt2, timedelta as _td2
                                    _start_dt = _dt2.fromisoformat(start_time.replace("Z", "+00:00"))
                                    end_time = (_start_dt + _td2(minutes=duracion_pedida)).isoformat()
                                except Exception as _dur_err:
                                    logger.error(f"Error calculando end_time por duracion de servicio: {_dur_err}")
                            if start_time and end_time:
                                # Misma verificacion de choque de horario que en el motor
                                # real (main.py): no confiar solo en el razonamiento del
                                # modelo, verificar a nivel de codigo antes de crear.
                                conflict_check = list_google_calendar_events(active_token, start_time, end_time)
                                if conflict_check.get("success") and conflict_check.get("busy_slots"):
                                    tool_result = json.dumps({
                                        "error": "Ese horario ya esta ocupado por otra cita existente. NO lo agendes. Ofrecele al cliente otro horario disponible.",
                                        "conflicto": True,
                                        "eventos_existentes": conflict_check["busy_slots"]
                                    })
                                else:
                                    res_event = create_google_calendar_event(
                                        active_token, summary, start_time, end_time, attendee_email, description, create_meet=cal_google_meet
                                    )
                                    tool_result = json.dumps(res_event)
                                    if res_event.get("success") and res_event.get("event_id"):
                                        # El simulador no tiene un contacto real en BD para guardar
                                        # el event_id, asi que se guarda igual que nombre/email: como
                                        # nota "Datos extraidos" en el historial, para poder
                                        # reconstruirlo despues si el cliente pide reagendar. Se agrega
                                        # directo a "notes" (no a res_data, que se reemplaza en la
                                        # siguiente iteracion del loop y perderia este dato).
                                        notes.append(f"📥 *[Simulación de Captura]* Datos extraídos: cita_event_id: *{res_event['event_id']}*")
                            else:
                                tool_result = '{"error": "Faltan parametros start_time o end_time (o start_time + servicio valido)"}'
                        elif tool_name == "reschedule_google_calendar_event":
                            new_start_time = tool_args.get("start_time")
                            servicio_pedido = _resolver_servicio_pedido(tool_args)
                            duracion_pedida = _lookup_duracion_servicio(servicio_pedido)
                            existing_event_id = None
                            for _h in reversed(history):
                                _h_text = _h.get('text') or ''
                                if 'cita_event_id' in _h_text:
                                    _m = re.search(r'cita_event_id:\s*\*([^*]*)\*', _h_text)
                                    if _m and _m.group(1).strip():
                                        existing_event_id = _m.group(1).strip()
                                        break
                            if not existing_event_id:
                                tool_result = '{"error": "No encontre ninguna cita previa de este cliente para reagendar en esta simulacion. Si quiere una cita nueva, usa create_google_calendar_event en su lugar."}'
                            elif not new_start_time:
                                tool_result = '{"error": "Falta el parametro start_time con la nueva hora deseada."}'
                            else:
                                new_end_time = tool_args.get("end_time")
                                if not new_end_time and duracion_pedida:
                                    try:
                                        from datetime import datetime as _dt3, timedelta as _td3
                                        _start_dt3 = _dt3.fromisoformat(new_start_time.replace("Z", "+00:00"))
                                        new_end_time = (_start_dt3 + _td3(minutes=duracion_pedida)).isoformat()
                                    except Exception as _dur_err2:
                                        logger.error(f"Error calculando end_time al reagendar en simulador: {_dur_err2}")
                                if not new_end_time:
                                    tool_result = '{"error": "No se pudo determinar la duracion de la cita para reagendar. Indica el servicio."}'
                                else:
                                    conflict_check = list_google_calendar_events(active_token, new_start_time, new_end_time)
                                    otros_eventos = [b for b in (conflict_check.get("busy_slots") or []) if b.get("id") != existing_event_id]
                                    if conflict_check.get("success") and otros_eventos:
                                        tool_result = json.dumps({
                                            "error": "Ese nuevo horario ya esta ocupado por otra cita. NO lo reagendes ahi. Ofrecele al cliente otro horario disponible.",
                                            "conflicto": True,
                                            "eventos_existentes": otros_eventos
                                        })
                                    else:
                                        from calendar_tools import update_google_calendar_event
                                        res_reschedule = update_google_calendar_event(active_token, existing_event_id, new_start_time, new_end_time)
                                        tool_result = json.dumps(res_reschedule)
                        elif tool_name == "cancel_google_calendar_event":
                            existing_event_id_cancel = None
                            for _h in reversed(history):
                                _h_text = _h.get('text') or ''
                                if 'cita_event_id' in _h_text:
                                    _m = re.search(r'cita_event_id:\s*\*([^*]*)\*', _h_text)
                                    if _m and _m.group(1).strip():
                                        existing_event_id_cancel = _m.group(1).strip()
                                        break
                            if not existing_event_id_cancel:
                                tool_result = '{"error": "No encontre ninguna cita previa de este cliente para cancelar en esta simulacion."}'
                            else:
                                from calendar_tools import delete_google_calendar_event
                                res_cancel = delete_google_calendar_event(active_token, existing_event_id_cancel)
                                tool_result = json.dumps(res_cancel)
                        else:
                            tool_result = f'{{"error": "Herramienta {tool_name} no valida."}}'
                    else:
                        tool_result = '{"error": "No se pudo obtener token de acceso de Google Calendar."}'
                except Exception as ex_tool:
                    logger.error(f"Error ejecutando herramienta de calendario en simulador: {ex_tool}")
                    tool_result = f'{{"error": "Error interno al ejecutar herramienta: {str(ex_tool)}"}}'
            elif cal_provider == "calendly" and cal_calendly_connected:
                try:
                    from calendar_tools import refresh_calendly_oauth_token, list_calendly_event_types
                    active_token = refresh_calendly_oauth_token(cursor, conn, agent["id"], config_json)
                    if active_token:
                        if tool_name == "list_calendly_slots":
                            res_events = list_calendly_event_types(active_token)
                            tool_result = json.dumps(res_events)
                        else:
                            tool_result = f'{{"error": "Herramienta {tool_name} no valida para Calendly."}}'
                    else:
                        tool_result = '{"error": "No se pudo obtener token de acceso de Calendly."}'
                except Exception as ex_tool:
                    logger.error(f"Error ejecutando herramienta de Calendly en simulador: {ex_tool}")
                    tool_result = f'{{"error": "Error interno al ejecutar herramienta Calendly: {str(ex_tool)}"}}'
            elif cal_provider == "cal.com" and cal_com_connected:
                try:
                    from calendar_tools import list_calcom_event_types
                    cal_api_key = config_json.get("calComApiKey")
                    cal_event_id = config_json.get("calComEventId")
                    if cal_api_key:
                        if tool_name == "list_calcom_slots":
                            res_events = list_calcom_event_types(cal_api_key, cal_event_id)
                            tool_result = json.dumps(res_events)
                        else:
                            tool_result = f'{{"error": "Herramienta {tool_name} no valida para Cal.com."}}'
                    else:
                        tool_result = '{"error": "API Key de Cal.com no configurada."}'
                except Exception as ex_tool:
                    logger.error(f"Error ejecutando herramienta de Cal.com en simulador: {ex_tool}")
                    tool_result = f'{{"error": "Error interno al ejecutar herramienta Cal.com: {str(ex_tool)}"}}'
            else:
                tool_result = f'{{"error": "Calendario no conectado en simulacion."}}'
                
            notes.append(f"🛠️ *[Simulacion de Herramienta]* Ejecutada: *{tool_name}*. Resultado: {tool_result[:150]}...")
            tool_results_context += f"- Herramienta: {tool_name}\n  Argumentos: {json.dumps(tool_args)}\n  Resultado: {tool_result}\n\n"
            logger.info(f"[TOOL_RESULT simulador] {tool_name} args={json.dumps(tool_args)} -> {tool_result}")

        tags_a_aplicar_ids = res_data.get("tags_a_aplicar_ids") or []
        regla_transferencia_id = res_data.get("regla_transferencia_id")
        datos_extraidos = res_data.get("datos_extraidos") or {}
        respuesta_final = (res_data.get("respuesta_final") or "").strip()
        seguimiento_data = res_data.get("seguimiento") or {}
        url_media_a_enviar = res_data.get("url_media_a_enviar") or None
        tipo_media_a_enviar = res_data.get("tipo_media_a_enviar") or None

        # Red de seguridad: el modelo a veces "se pierde" en conversaciones largas y
        # responde sin extraer el dato que el cliente acaba de dar (ej. reinicia el
        # saludo en vez de continuar). Si el paso pendiente era un correo y el ultimo
        # mensaje del cliente tiene forma de email, lo guardamos de todas formas aunque
        # la IA no lo haya detectado — no depender solo de que el modelo lo razone bien.
        if pending_field_name in ("email", "correo") and not (datos_extraidos.get("email") or datos_extraidos.get("correo")):
            email_match = re.search(r'[\w\.\-+]+@[\w\-]+\.[\w\.\-]+', message_text or "")
            if email_match:
                datos_extraidos["email"] = email_match.group(0)
                logger.warning(f"Agente {agent.get('nombre')} (simulador): la IA no extrajo el correo, se recupero con regex de respaldo: {email_match.group(0)}")

        # Red de seguridad CRITICA: verificar que el correo que la IA dice haber
        # "extraido" en datos_extraidos realmente aparece en el mensaje del cliente,
        # y no fue inventado a partir del nombre.
        _correo_declarado = datos_extraidos.get("correo") or datos_extraidos.get("email")
        if _correo_declarado:
            if _correo_declarado.strip().lower() not in (message_text or "").lower():
                logger.warning(f"Agente {agent.get('nombre')} (simulador): la IA reporto un correo ('{_correo_declarado}') que no aparece en el mensaje del cliente; se descarta por posible invencion.")
                datos_extraidos.pop("correo", None)
                datos_extraidos.pop("email", None)

        # Red de seguridad: si se agotaron los intentos sin que el modelo llegara a dar
        # una respuesta final (se quedo pidiendo la misma herramienta una y otra vez),
        # no dejar al cliente sin ningun mensaje.
        if not respuesta_final and res_data.get("tool_call"):
            respuesta_final = "Estoy verificando la disponibilidad, dame un momento y te confirmo por aquí mismo."
            logger.warning(f"Agente {agent.get('nombre')}: se agotaron los intentos sin respuesta final, se uso mensaje de respaldo.")

        # Red de seguridad CRITICA: nunca dejar que el cliente crea que su cita quedo
        # agendada/cancelada/reagendada si en realidad no se ejecuto con exito la
        # herramienta correspondiente. Se confirmo en pruebas reales que la IA a
        # veces redacta una confirmacion sin haber llamado a la herramienta (o
        # habiendola llamado sin exito). Se verifica el resultado real de la
        # herramienta, no lo que diga el texto.
        if respuesta_final and agent.get("objetivo") == "agendar_citas":
            _fake_confirm_checks = [
                (
                    r'cita.{0,40}(ha sido|qued[oó]|est[aá])\s+confirmad|confirmad[oa].{0,40}\bcita\b',
                    "create_google_calendar_event",
                    "Estoy verificando la disponibilidad real para ese horario antes de confirmarte — "
                    "dame un momento y te aviso apenas quede agendada la cita.",
                ),
                (
                    r'cita.{0,40}(ha sido|qued[oó]|est[aá])\s+cancelad|cancelad[oa].{0,40}\bcita\b',
                    "cancel_google_calendar_event",
                    "Estoy procesando la cancelación de tu cita — dame un momento y te confirmo apenas quede cancelada.",
                ),
                (
                    r'cita.{0,40}(ha sido|qued[oó]|est[aá])\s+reprogramad|reprogramad[oa].{0,40}\bcita\b|cita.{0,40}(ha sido|qued[oó])\s+movid|cambiad[oa].{0,40}\bcita\b',
                    "reschedule_google_calendar_event",
                    "Estoy verificando ese nuevo horario antes de confirmarte el cambio — dame un momento y te aviso.",
                ),
            ]
            for _pattern, _tool, _fallback in _fake_confirm_checks:
                if not re.search(_pattern, respuesta_final, re.IGNORECASE):
                    continue
                _tool_ok = bool(re.search(
                    r'Herramienta:\s*' + re.escape(_tool) + r'\b.*?Resultado:\s*\{[^\n]*"success":\s*true',
                    tool_results_context, re.IGNORECASE | re.DOTALL
                ))
                if not _tool_ok:
                    logger.warning(f"Agente {agent.get('nombre')}: la respuesta afirmaba una accion de '{_tool}' SIN una llamada exitosa a esa herramienta; se reemplaza para no engañar al cliente. Respuesta original: {respuesta_final}")
                    respuesta_final = _fallback
                break

        # Validar que la URL de media pertenezca a un recurso real del agente (seguridad)
        if url_media_a_enviar and recursos_rows:
            from urllib.parse import urlparse
            def _extract_path(url):
                """Extrae la parte de path de una URL o devuelve la cadena tal cual si no es URL."""
                try:
                    p = urlparse(url)
                    return p.path.lstrip('/')
                except Exception:
                    return url or ''

            path_enviada = _extract_path(url_media_a_enviar)
            # Buscar el recurso cuyo path coincida con el enviado por la IA
            recurso_match = next(
                (r for r in recursos_rows if _extract_path(r.get('archivo_url', '')) == path_enviada),
                None
            )
            if recurso_match:
                # Usar la URL canónica almacenada en la BD (para consistencia)
                url_media_a_enviar = recurso_match['archivo_url']
            else:
                logger.warning(f"IA devolvió URL de media inválida (path no coincide): {url_media_a_enviar}")
                url_media_a_enviar = None
                tipo_media_a_enviar = None

        # Segunda capa: si la IA igual coló la URL en el texto, la removemos
        if url_media_a_enviar:
            import re as _re
            # Quitar la URL exacta o cualquier URL que contenga el path del recurso
            respuesta_final = respuesta_final.replace(url_media_a_enviar, '').strip()
            # Quitar cualquier otra URL http/https que quedara (por si la IA usó otro dominio)
            respuesta_final = _re.sub(r'https?://\S+', '', respuesta_final).strip()
            # Limpiar frases residuales como "en el siguiente enlace:" al final o al medio
            respuesta_final = _re.sub(r'\s*(en el siguiente enlace|a través del enlace|aquí el enlace|link)\s*:?\s*', ' ', respuesta_final, flags=_re.IGNORECASE).strip()
            respuesta_final = _re.sub(r'\s{2,}', ' ', respuesta_final).strip()

        notes = []
        if parsed_ok:
            if tags_a_aplicar_ids and reglas_etiquetado_raw:
                try:
                    reglas_etiquetado = json.loads(reglas_etiquetado_raw)
                    matched_ids_str = [str(x) for x in tags_a_aplicar_ids]
                    applied_rules = []
                    for rule in reglas_etiquetado:
                        if str(rule.get("id")) in matched_ids_str:
                            action_type = rule.get('action') or rule.get('type') or 'Agregar'
                            label_name = rule.get('label') or rule.get('target') or 'Sin etiqueta'
                            applied_rules.append(f"{action_type} etiqueta *{label_name}*")
                    if applied_rules:
                        notes.append(f"🏷️ *[Simulación de Etiquetas]* Se ejecutó: {', '.join(applied_rules)}")
                except Exception as tag_err:
                    logger.error(f"Error en notas de etiquetas de simulación: {tag_err}")

            if regla_transferencia_id is not None and reglas_trans_raw:
                try:
                    reglas_trans = json.loads(reglas_trans_raw)
                    matched_rule = next((r for r in reglas_trans if str(r.get("id")) == str(regla_transferencia_id)), None)
                    if matched_rule:
                        target = matched_rule.get("target")
                        dest_type = matched_rule.get("type")
                        if dest_type == "Humano":
                            notes.append(f"🔄 *[Simulación de Transferencia]* Derivado a asesor humano: *{target}*")
                        elif dest_type == "Superagente":
                            notes.append(f"🔄 *[Simulación de Transferencia]* Derivado a superagente virtual: *{target}*")
                        elif dest_type == "Flujo":
                            notes.append(f"🔄 *[Simulación de Transferencia]* Transferido a automatización de flujo: *{target}*")
                except Exception as trans_err:
                    logger.error(f"Error en notas de transferencia de simulación: {trans_err}")

            if datos_extraidos:
                extracted_notes = [f"{k}: *{v}*" for k, v in datos_extraidos.items() if v]
                if extracted_notes:
                    notes.append(f"📥 *[Simulación de Captura]* Datos extraídos: {', '.join(extracted_notes)}")

            if seguimiento_enabled and seguimiento_data.get("programar") is True:
                horas = seguimiento_data.get("horas_retraso") or 24
                msg = seguimiento_data.get("mensaje_propuesto") or ""
                notes.append(f"⏰ *[Simulación de Seguimiento]* Se programará recordatorio en {horas} horas: \"{msg}\"")

        if not respuesta_final:
            detail_err = "No se pudo generar respuesta. Detalle del estado de APIs:\n"
            if errors_list:
                detail_err += "\n".join(f"- {err}" for err in errors_list)
            else:
                detail_err += "- Proveedores de IA deshabilitados o sin keys."
            respuesta_final = f"⚠️ {detail_err}"

        return jsonify({
            "success": True,
            "reply": respuesta_final,
            "notes": notes,
            "transcription": transcription,
            "url_media": url_media_a_enviar,
            "tipo_media": tipo_media_a_enviar
        })

    except Exception as e:
        logger.exception("Error al simular mensaje del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/audit', methods=['POST'])
@jwt_required()
def audit_agent_config(agent_id):
    user_id = get_jwt_identity()
    payload = request.get_json() or {}
    action = payload.get('action', 'analizar')
    message_text = payload.get('message', '').strip()
    history = payload.get('history', [])
    
    openai_key = os.getenv("OPENAI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    nvidia_key = os.getenv("NVIDIA_API_KEY")
    
    if not openai_key and not gemini_key and not nvidia_key:
        return jsonify({
            "success": True, 
            "reply": "⚠️ No hay API keys configuradas en el servidor de procesos para interactuar con la auditoría."
        })

    conn = None
    cursor = None
    try:
        from main import call_llm_api
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT * FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Agente no encontrado"}), 404

        cursor.execute("SELECT titulo FROM agente_conocimiento WHERE agente_id = %s", (agent_id,))
        conocimiento_list = [row['titulo'] for row in cursor.fetchall()]

        # Parsear todos los bloques de configuración para darle al auditor visión
        # completa del agente (antes solo veía Reglas, no Rol, Negocio, Pasos,
        # Calendario, Voz, etc. — dejaba pasar cosas importantes).
        def _safe_json_list(raw):
            if not raw:
                return []
            try:
                parsed = json.loads(raw)
                return parsed if isinstance(parsed, list) else []
            except Exception:
                return []

        reglas_trans_list = _safe_json_list(agent.get("reglas_transferencia"))
        reglas_etiq_list = _safe_json_list(agent.get("reglas_etiquetado"))
        pasos_list = _safe_json_list(agent.get("pasos_captura"))
        seguimientos_list = _safe_json_list(agent.get("seguimientos"))

        config_comp = {}
        if agent.get("config_comportamiento"):
            try:
                config_comp = json.loads(agent["config_comportamiento"])
            except Exception:
                pass

        transferencias_text = "Ninguna configurada."
        if reglas_trans_list:
            transferencias_text = "; ".join(
                f"[{r.get('type')}] si \"{r.get('text')}\" → destino: \"{r.get('target') or 'SIN DESTINO'}\""
                for r in reglas_trans_list
            )

        etiquetado_text = "Ninguna configurada."
        if reglas_etiq_list:
            etiquetado_text = "; ".join(
                f"si \"{r.get('text')}\" → {r.get('action') or 'Agregar'} etiqueta \"{r.get('label') or r.get('target')}\""
                for r in reglas_etiq_list
            )

        pasos_text_audit = "Ninguno configurado."
        if pasos_list:
            pasos_text_audit = "; ".join(
                f"{'[ACTIVO]' if p.get('enabled') else '[INACTIVO]'} {p.get('text')} (campo: {p.get('field') or 'no guarda dato'})"
                for p in pasos_list
            )

        seguimientos_text = "Ninguno configurado."
        if seguimientos_list:
            seguimientos_text = "; ".join(
                f"\"{s.get('text')}\" a los {s.get('time')} {s.get('unit')}"
                for s in seguimientos_list
            )

        cal_provider = config_comp.get("calProvider")
        cal_connected = bool(
            config_comp.get("calGoogleConnected") or config_comp.get("calCalendlyConnected") or config_comp.get("calComApiKey")
        )

        config_status = {
            "nombre": agent.get("nombre"),
            "giro": agent.get("industria"),
            "objetivo": agent.get("objetivo"),
            "personalidad": agent.get("personalidad"),
            "descripcion_negocio": agent.get("descripcion_negocio"),
            "instrucciones": agent.get("instrucciones"),
            "tiene_transferencia": bool(reglas_trans_list),
            "tiene_etiquetado": bool(reglas_etiq_list),
            "conocimiento_docs": conocimiento_list
        }

        history_text = ""
        for h in history[-8:]:
            sender_label = "Usuario" if h.get("sender") == "user" else "Asistente"
            history_text += f"{sender_label}: {h.get('text')}\n"
        if message_text:
            history_text += f"Usuario: {message_text}\n"

        system_prompt = (
            "Eres el Asistente de Configuración (Auditor de IAs) de la plataforma GeoCHAT.\n"
            "Tu labor es auditar y dar sugerencias de optimización personalizadas, concretas y accionables para el "
            "superagente del usuario en español. Sé específico: si algo está vacío, mal configurado, incompleto o "
            "podría hacer que el bot falle en atender bien a un cliente real, dilo explícitamente y explica cómo corregirlo.\n"
            f"El agente actual tiene la siguiente configuración:\n"
            f"- Nombre: {config_status['nombre']}\n"
            f"- Giro del Negocio: {config_status['giro']}\n"
            f"- Objetivo: {config_status['objetivo']}\n"
            f"- Rol (personalidad): {config_status['personalidad'] or 'VACÍO'}\n"
            f"- Información del negocio: {config_status['descripcion_negocio'] or 'VACÍO'}\n"
            f"- Reglas de conversación (instrucciones): {config_status['instrucciones'] or 'VACÍO'}\n"
            f"- Pasos de captura de datos: {pasos_text_audit}\n"
            f"- Reglas de transferencia a humano: {transferencias_text}\n"
            f"- Reglas de etiquetado automático: {etiquetado_text}\n"
            f"- Seguimientos automáticos configurados: {seguimientos_text}\n"
            f"- Seguimiento Inteligente (detecta 'escríbeme luego' y agenda solo): {'Activado' if config_comp.get('seguimientoInteligente') else 'Desactivado'}\n"
            f"- Calendario: proveedor '{cal_provider or 'ninguno'}', {'conectado' if cal_connected else 'NO conectado'}"
            f"{' (pero el objetivo no es Agendar Citas, así que el calendario está inactivo aunque esté conectado)' if cal_connected and config_status['objetivo'] != 'agendar_citas' else ''}"
            f"{' (el objetivo es Agendar Citas pero no hay calendario conectado, el bot no podrá agendar de verdad)' if config_status['objetivo'] == 'agendar_citas' and not cal_connected else ''}\n"
            f"- Respuestas de voz: {'Activadas al ' + str(config_comp.get('voicePercentage', 0)) + '%' if config_comp.get('voiceEnabled') else 'Desactivadas'}\n"
            f"- Documentos de conocimiento cargados: {', '.join(config_status['conocimiento_docs']) if config_status['conocimiento_docs'] else 'Ninguno'}\n\n"
            "El usuario ha solicitado la siguiente acción de auditoría: "
        )
        
        if action == 'resolver':
            import time
            generator_prompt = (
                "Eres un optimizador de agentes de IA para GeoCHAT.\n"
                f"El agente '{agent.get('nombre')}' tiene el giro '{agent.get('industria')}' y el objetivo '{agent.get('objetivo')}'.\n"
                f"Su descripción actual es: \"{agent.get('descripcion_negocio') or ''}\"\n\n"
                "Genera un objeto JSON con las siguientes propiedades (estrictamente JSON, no incluyes markdown fuera del bloque JSON, ni explicaciones):\n"
                "{\n"
                "  \"instrucciones\": \"Un prompt de comportamiento profesional completo y detallado para este agente (al menos 3 oraciones).\",\n"
                "  \"transfer_rule_condition\": \"Una condición clara de una sola oración para transferir a un humano (ej: cuando pida hablar con un doctor o tenga urgencia).\",\n"
                "  \"follow_up_message\": \"Un mensaje cálido de recordatorio si el cliente deja de responder.\"\n"
                "}"
            )
            llm_res = call_llm_api(generator_prompt, "Generador de resolución", openai_key, gemini_key, nvidia_key)
            
            optimized_instrucciones = None
            transfer_cond = "Cuando el cliente tenga una duda compleja o pida hablar con un humano."
            follow_up_msg = f"Hola, soy {agent.get('nombre')}. ¿Sigues por ahí? Avísame si tienes alguna duda."
            
            try:
                import re
                json_match = re.search(r'\{.*\}', llm_res.strip(), re.DOTALL)
                if json_match:
                    data_parsed = json.loads(json_match.group(0))
                    optimized_instrucciones = data_parsed.get("instrucciones")
                    transfer_cond = data_parsed.get("transfer_rule_condition") or transfer_cond
                    follow_up_msg = data_parsed.get("follow_up_message") or follow_up_msg
            except Exception as pe:
                logger.error(f"Error parsing resolver JSON: {pe}")
                
            if not optimized_instrucciones:
                optimized_instrucciones = (
                    f"Eres el asistente de {agent.get('nombre') or 'DentiSmile'}. Ayudas a los clientes a resolver "
                    f"dudas sobre el negocio y agendar citas. Responde de forma cálida, profesional y concisa."
                )

            # Obtener el nombre del usuario actual (dueño) para asignar la transferencia de forma válida
            cursor.execute("SELECT nombre FROM usuarios WHERE id = %s LIMIT 1", (user_id,))
            user_row = cursor.fetchone()
            owner_name = user_row["nombre"] if (user_row and user_row.get("nombre")) else "Soporte"

            existing_trans = []
            if agent.get("reglas_transferencia"):
                try:
                    existing_trans = json.loads(agent["reglas_transferencia"])
                except:
                    pass

            # Si ya hay reglas sin destino asignado (el gap "Regla de transferencia sin
            # destino"), reemplázalas en vez de apilar una regla nueva encima y dejar la
            # rota ahí — si no, el mismo gap seguiría apareciendo en la próxima auditoría.
            existing_trans = [r for r in existing_trans if (r.get("target") or "").strip() and r.get("target") != "Elegir..."]

            new_rule = {
                "id": int(time.time() * 1000),
                "text": transfer_cond,
                "type": "Humano",
                "target": owner_name
            }
            existing_trans.append(new_rule)
            
            existing_comp = {}
            if agent.get("config_comportamiento"):
                try:
                    existing_comp = json.loads(agent["config_comportamiento"])
                except:
                    pass
            existing_comp["seguimientoInteligente"] = True
            
            existing_seguimientos = []
            if agent.get("seguimientos"):
                try:
                    existing_seguimientos = json.loads(agent["seguimientos"])
                except:
                    pass
            new_followup = {
                "id": int((time.time() * 1000) + 1),
                "text": follow_up_msg,
                "time": 30,
                "unit": "min"
            }
            existing_seguimientos.append(new_followup)
            
            cursor.execute("""
                UPDATE agentes_ia 
                SET instrucciones = %s, 
                    reglas_transferencia = %s, 
                    config_comportamiento = %s,
                    seguimientos = %s
                WHERE id = %s AND usuario_id = %s
            """, (
                optimized_instrucciones,
                json.dumps(existing_trans),
                json.dumps(existing_comp),
                json.dumps(existing_seguimientos),
                agent_id,
                user_id
            ))
            conn.commit()
            
            reply = (
                "### ¡Cambios Aplicados Automáticamente! 🎉\n\n"
                f"He optimizado la configuración de tu superagente **{agent.get('nombre')}** directamente en la base de datos:\n\n"
                f"1. **Instrucciones de comportamiento actualizadas**:\n"
                f"   > *\"{optimized_instrucciones}\"*\n\n"
                f"2. **Regla de transferencia creada**:\n"
                f"   * Condición: *\"{transfer_cond}\"*\n"
                f"   * Acción: Derivar al asesor humano.\n\n"
                f"3. **Seguimiento inteligente activado**:\n"
                f"   * Mensaje de seguimiento: *\"{follow_up_msg}\"*\n"
                f"   * Tiempo de espera: 30 minutos de inactividad.\n\n"
            )

            cursor.execute("SELECT COUNT(*) as count FROM agente_conocimiento WHERE agente_id = %s", (agent_id,))
            if cursor.fetchone()["count"] == 0:
                reply += (
                    "⚠️ Nota: si la 'Base de conocimiento vacía' te salió como problema, eso **no lo puedo resolver automáticamente** — "
                    "necesito que subas tú misma documentos, FAQs o páginas web en la pestaña 'Conocimiento'.\n\n"
                )

            reply += "**Por favor, refresca la página (F5 o vuelve a ingresar al agente) para ver estos cambios reflejados en tus pestañas de configuración.**"
            return jsonify({
                "success": True,
                "reply": reply
            })

        if action == 'analizar':
            system_prompt += "Escribe un reporte de auditoría completo y profesional, estructurado en secciones cortas (Gaps, Sugerencias, Aspectos Correctos) adaptadas a su giro de negocio."
        elif action == 'instrucciones':
            system_prompt += "Propón una optimización del texto de sus instrucciones (prompt de comportamiento) para que el bot responda de forma más asertiva y natural."
        elif action == 'mejoras':
            system_prompt += "Proporciona 3 ideas avanzadas de negocio para expandir las capacidades y el valor de su asistente."
        elif action == 'chat':
            system_prompt += "Responde al último mensaje del usuario de manera directa, conversacional y amigable. Si saluda, salúdalo y ofrécete a ayudar. Si hace preguntas, respóndelas directamente sin volver a generar el reporte completo de auditoría."
            
        system_prompt += (
            "\n\nHistorial de interacción con el Asistente de Configuración:\n"
            f"{history_text}\n"
        )
        
        if action == 'chat':
            system_prompt += "Responde directamente en español al último mensaje del usuario de forma conversacional y constructiva."
        else:
            system_prompt += "Escribe tu reporte de auditoría directamente en español. Usa un tono constructivo, profesional, claro y de alto valor comercial."

        response_text = call_llm_api(system_prompt, f"Auditor de Configuración - {agent.get('nombre')}", openai_key, gemini_key, nvidia_key)
        
        return jsonify({
            "success": True,
            "reply": (response_text or "").strip()
        })
    except Exception as e:
        logger.exception("Error al auditar configuración del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@agentes_ia_blueprint.route('/api/agentes-ia/<int:agent_id>/audit-status', methods=['GET'])
@jwt_required()
def audit_agent_status(agent_id):
    user_id = get_jwt_identity()
    conn = None
    cursor = None
    try:
        conn = get_connection()
        cursor = conn.cursor(dictionary=True)
        
        cursor.execute("SELECT * FROM agentes_ia WHERE id = %s AND usuario_id = %s", (agent_id, user_id))
        agent = cursor.fetchone()
        if not agent:
            return jsonify({"success": False, "message": "Agente no encontrado"}), 404
            
        cursor.execute("SELECT count(*) as count FROM agente_conocimiento WHERE agente_id = %s", (agent_id,))
        conocimiento_count = cursor.fetchone()['count']
        
        gaps = []
        
        # 1. Reglas de transferencia
        reglas_trans = []
        if agent.get("reglas_transferencia"):
            try:
                reglas_trans = json.loads(agent["reglas_transferencia"])
            except:
                pass
        if not reglas_trans:
            gaps.append({
                "type": "Faltante",
                "title": "Reglas de transferencia",
                "detail": "No hay reglas de transferencia configuradas. Si el agente no puede resolver una duda compleja, no podrá derivar la conversación a un asesor humano."
            })
        else:
            for rule in reglas_trans:
                rule_target = (rule.get("target") or "").strip()
                if not rule_target or rule_target == "Elegir...":
                    gaps.append({
                        "type": "Faltante",
                        "title": "Regla de transferencia sin destino",
                        "detail": f"La regla \"{rule.get('text') or 'sin condición'}\" no tiene un destino seleccionado (sigue en 'Elegir...'). Mientras no elijas a quién transferir, esta regla no hace nada aunque se cumpla la condición."
                    })

        # 1.b Reglas de etiquetado con condición sin definir
        reglas_etiq = []
        if agent.get("reglas_etiquetado"):
            try:
                reglas_etiq = json.loads(agent["reglas_etiquetado"])
            except:
                pass
        for rule in reglas_etiq:
            rule_text = (rule.get("text") or "").strip()
            if not rule_text or rule_text.lower() == "nueva condición":
                gaps.append({
                    "type": "Faltante",
                    "title": "Regla de etiquetado sin condición",
                    "detail": f"Hay una regla que agrega la etiqueta \"{rule.get('label') or rule.get('target') or 'sin nombre'}\" pero su condición sigue siendo el texto de ejemplo (\"Nueva condición\") o está vacía. Así, la IA puede aplicarla de forma imprecisa en casi cualquier mensaje. Describe una condición específica y real."
                })

        # 2. Seguimientos automáticos
        config_comp = {}
        if agent.get("config_comportamiento"):
            try:
                config_comp = json.loads(agent["config_comportamiento"])
            except:
                pass
        
        if not config_comp.get("seguimientoInteligente"):
            gaps.append({
                "type": "Recomendado",
                "title": "Seguimiento Inteligente inactivo",
                "detail": "El interruptor de 'Seguimiento Inteligente' está desactivado en la pestaña Auto-Tareas. El bot no programará recordatorios si el cliente deja de responder."
            })
            
        # 3. Base de conocimiento
        if conocimiento_count == 0:
            gaps.append({
                "type": "Faltante",
                "title": "Base de conocimiento vacía",
                "detail": "No has cargado documentos ni FAQs en la pestaña 'Conocimiento'. El agente solo responderá con sus instrucciones generales y podría inventar información."
            })
            
        # 4. Instrucciones generales
        descripcion = (agent.get("descripcion_negocio") or "").strip()
        if len(descripcion) < 15:
            gaps.append({
                "type": "Mejora",
                "title": "Descripción del negocio muy breve",
                "detail": f"La descripción del negocio es demasiado corta (actualmente: '{descripcion}'). Agrega más detalles sobre tus servicios, horarios o dirección para entrenar mejor al bot."
            })

        # 4.b Rol / Personalidad
        personalidad = (agent.get("personalidad") or "").strip()
        if len(personalidad) < 15:
            gaps.append({
                "type": "Mejora",
                "title": "Rol del asistente muy breve",
                "detail": "El texto de 'Rol' (pestaña General → Editar Instrucciones) es muy corto o está vacío. Sin un rol claro, el bot puede sonar genérico o inconsistente con el tono de tu negocio."
            })

        # 5. Pasos de captura
        pasos = []
        if agent.get("pasos_captura"):
            try:
                pasos = json.loads(agent["pasos_captura"])
            except:
                pass
        if not pasos:
            gaps.append({
                "type": "Opcional",
                "title": "Sin pasos de captura de datos",
                "detail": "No hay campos de captura (como Nombre o Correo) configurados en la pestaña 'Conversación'. El bot no almacenará datos estructurados del cliente."
            })

        # 6. Calendario: conectado pero inactivo, u objetivo sin calendario
        cal_connected = bool(
            config_comp.get("calGoogleConnected") or config_comp.get("calCalendlyConnected") or config_comp.get("calComApiKey")
        )
        objetivo_actual = agent.get("objetivo")
        if cal_connected and objetivo_actual != "agendar_citas":
            gaps.append({
                "type": "Mejora",
                "title": "Calendario conectado pero inactivo",
                "detail": "Conectaste un calendario (Google/Calendly/Cal.com), pero el objetivo actual del agente no es 'Agendar Citas', así que el bot no usará el calendario aunque siga conectado. Si ya no lo necesitas para este agente, puedes desconectarlo; si sí, cambia el objetivo a 'Agendar Citas'."
            })
        elif objetivo_actual == "agendar_citas" and not cal_connected:
            gaps.append({
                "type": "Faltante",
                "title": "Agenda de citas sin calendario conectado",
                "detail": "El objetivo del agente es 'Agendar Citas', pero no hay ningún calendario (Google/Calendly/Cal.com) conectado en la pestaña Conversación → Calendario. El bot no podrá agendar citas reales todavía."
            })

        # 7. Respuestas de voz activadas al 0%
        if config_comp.get("voiceEnabled") and int(config_comp.get("voicePercentage") or 0) == 0:
            gaps.append({
                "type": "Opcional",
                "title": "Respuestas de voz activadas pero al 0%",
                "detail": "Tienes 'Respuestas de Voz' activado en la pestaña Voz, pero el porcentaje está en 0% — el bot nunca va a responder con audio en la práctica. Sube el porcentaje o desactiva la opción."
            })

        return jsonify({
            "success": True,
            "gaps": gaps,
            "problems_count": len(gaps)
        })
    except Exception as e:
        logger.exception("Error al obtener estado de auditoría del agente")
        return jsonify({"success": False, "message": str(e)}), 500
    finally:
        if cursor:
            cursor.close()
        if conn:
            conn.close()


@agentes_ia_blueprint.route('/api/agentes-ia/optimize-prompt', methods=['POST'])
@jwt_required()
def optimize_agent_prompt():
    user_id = get_jwt_identity()
    payload = request.get_json() or {}
    text_draft = payload.get('draft', '').strip()
    field_type = payload.get('type', 'rol') # rol, negocio, reglas
    
    if not text_draft:
        return jsonify({"success": False, "message": "El borrador de texto está vacío"}), 400
        
    openai_key = os.getenv("OPENAI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    nvidia_key = os.getenv("NVIDIA_API_KEY")
    
    if not openai_key and not gemini_key and not nvidia_key:
        return jsonify({
            "success": True, 
            "optimized": text_draft,
            "message": "⚠️ No hay API keys configuradas en el servidor para realizar la optimización."
        })
        
    system_prompt = (
        "Eres un experto Ingeniero de Prompts especialista en modelos de lenguaje grandes (LLMs) y agentes conversacionales.\n"
        "El usuario te proporcionará un borrador rudimentario o simple para configurar una sección de su superagente de IA.\n"
        f"La sección que se está optimizando es: {field_type.upper()}.\n\n"
        "Instrucciones:\n"
        "1. Reescribe y optimiza el borrador del usuario para convertirlo en una instrucción profesional, estructurada y muy clara para que la IA la siga al pie de la letra.\n"
        "2. Mantén la esencia y todos los detalles del negocio o reglas que el usuario ingresó, pero redáctalo en un español impecable, usando un tono profesional, asertivo y de alto valor comercial.\n"
        "3. Responde ÚNICAMENTE con el texto optimizado. No agregues introducciones, explicaciones, ni saludos (ej: no digas 'Aquí tienes el prompt optimizado:'). Escribe directamente el prompt mejorado.\n\n"
        f"Borrador del usuario:\n\"{text_draft}\""
    )
    
    try:
        from main import call_llm_api
        optimized_text = call_llm_api(system_prompt, f"Optimizador de Prompt ({field_type})", openai_key, gemini_key, nvidia_key)
        return jsonify({
            "success": True,
            "optimized": (optimized_text or "").strip()
        })
    except Exception as e:
        logger.exception("Error al optimizar prompt")
        return jsonify({"success": False, "message": str(e)}), 500

@agentes_ia_blueprint.route('/api/agentes-ia/generate-message', methods=['POST'])
@jwt_required()
def generate_whatsapp_message():
    user_id = get_jwt_identity()
    payload = request.get_json() or {}
    draft = payload.get('draft', '').strip()
    instruction = payload.get('instruction', '').strip()
    action = payload.get('action', 'improve').strip()

    openai_key = os.getenv("OPENAI_API_KEY")
    gemini_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    nvidia_key = os.getenv("NVIDIA_API_KEY")

    if not openai_key and not gemini_key and not nvidia_key:
        return jsonify({
            "success": True,
            "generated": draft or "¡Hola! Gracias por contactarnos. ¿En qué podemos ayudarte hoy?",
            "message": "No hay API keys configuradas en el servidor. Se devolvió texto por defecto."
        })

    if action == 'generate' or instruction:
        system_prompt = (
            "Eres un redactor profesional experto en copywriting y marketing conversacional para WhatsApp.\n"
            "El usuario te ha pedido ayuda para redactar o mejorar un mensaje de difusión masiva para sus clientes.\n\n"
            f"Instrucción del usuario: \"{instruction}\"\n"
            f"Borrador o contexto actual: \"{draft}\"\n\n"
            "Instrucciones para la respuesta:\n"
            "1. Redacta un mensaje atractivo, persuasivo, claro y optimizado para WhatsApp (puedes usar negritas con *, saltos de línea y emojis si son apropiados).\n"
            "2. Mantén un tono natural, amigable y profesional. Evita sonar robótico.\n"
            "3. Responde ÚNICAMENTE con el cuerpo del mensaje generado. No agregues introducciones, explicaciones, ni textos como 'Aquí tienes tu mensaje:' o saludos iniciales para el usuario. Escribe directamente el mensaje final para el cliente.\n"
        )
    else:
        system_prompt = (
            "Eres un redactor profesional experto en copywriting y marketing conversacional para WhatsApp.\n"
            "El usuario te ha pedido optimizar y mejorar su borrador de mensaje de difusión masiva.\n\n"
            f"Borrador actual del usuario:\n\"{draft}\"\n\n"
            "Instrucciones para la respuesta:\n"
            "1. Optimiza la redacción, ortografía y fluidez del borrador. Hazlo más persuasivo y atractivo para los clientes.\n"
            "2. Mantén la esencia, información relevante y los marcadores de posición (por ejemplo, variables como {nombre} o {telefono}) si existen.\n"
            "3. Usa un formato limpio con saltos de línea adecuados y emojis oportunos.\n"
            "4. Responde ÚNICAMENTE con el mensaje optimizado. No agregues introducciones, explicaciones ni comentarios. Escribe directamente el texto final."
        )

    try:
        from main import call_llm_api
        generated_text = call_llm_api(system_prompt, f"Asistente Redacción AI ({action})", openai_key, gemini_key, nvidia_key)
        return jsonify({
            "success": True,
            "generated": (generated_text or "").strip()
        })
    except Exception as e:
        logger.exception("Error al generar mensaje con AI")
        return jsonify({"success": False, "message": str(e)}), 500

